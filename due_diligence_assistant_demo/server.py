from __future__ import annotations

import json
import hashlib
import mimetypes
import os
import re
import sqlite3
import ssl
from cgi import FieldStorage
from collections import defaultdict
from datetime import datetime
from functools import lru_cache
from html import escape as html_escape
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from random import Random
from urllib.parse import parse_qs, quote, unquote, urlparse
from urllib.request import Request, urlopen
from zipfile import ZIP_DEFLATED, ZipFile

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document
import certifi

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    REPORTLAB_AVAILABLE = True
except ImportError:  # pragma: no cover - optional runtime fallback
    REPORTLAB_AVAILABLE = False


ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = ROOT.parent
DB_PATH = PROJECT_ROOT / "database" / "sme_due_diligence_demo.db"
XLSX_PATH = PROJECT_ROOT / "outputs" / "sme_due_diligence_demo" / "sme_due_diligence_demo.xlsx"
LOCAL_CONFIG_PATH = ROOT / ".local_config.json"
KNOWLEDGE_LIBRARY_DIR = ROOT / "knowledge_library"
KNOWLEDGE_FILES_DIR = KNOWLEDGE_LIBRARY_DIR / "files"
KNOWLEDGE_META_PATH = KNOWLEDGE_LIBRARY_DIR / "manifest.json"
HOST = os.getenv("HOST", "0.0.0.0")
PORT = int(os.getenv("PORT") or os.getenv("DUE_DILIGENCE_DEMO_PORT", "8765"))
PDF_FONT_NAME = "ArialUnicodeMS"
PDF_FONT_PATHS = [
    Path("/Library/Fonts/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
]
REPORT_RUNTIME: dict[str, dict] = {}
DUE_DILIGENCE_MATERIALS_RUNTIME: dict[str, dict] = {}
KNOWLEDGE_BASE_RUNTIME: dict[str, object] = {"initialized": False, "next_id": 1, "files": []}
GLOBAL_SKILL_PIPELINE = {
    "label": "通用双 Skill 报告生成",
    "skills": ["auditing-data", "writing-report"],
    "entry_mode": "auditing_then_writing",
}
SKILL_POC_COMPANY_CODE = "COMP-001"
SKILL_PIPELINE_MOCK_BINDINGS = {
    SKILL_POC_COMPANY_CODE: {
        "mock_vdr_dir": PROJECT_ROOT / ".agent" / "mock_vdr" / "lingxi-sensors",
        "company_name": "深圳市灵犀微传感科技有限公司",
    }
}
MATERIAL_BUCKET_TITLES = {
    "application": "授信申请材料",
    "financial": "财务与经营材料",
    "governance": "主体与治理材料",
    "collateral": "担保与抵质押材料",
    "external": "外部核验材料",
    "manager-notes": "客户经理补充说明",
}
KNOWLEDGE_CATEGORY_TITLES = {
    "laws": "外部法规",
    "policies": "内部制度",
    "experience": "专家经验",
}
REPORT_DATA_SOURCE_GROUPS = [
    {
        "id": "customer-base",
        "title": "客户与案例基础",
        "items": [
            {"id": "company", "label": "企业主体信息"},
            {"id": "case_info", "label": "案例与授信申请信息"},
            {"id": "people", "label": "核心人员与角色信息"},
            {"id": "profile_attributes", "label": "补充画像属性"},
        ],
    },
    {
        "id": "operation-flow",
        "title": "经营与交易闭环",
        "items": [
            {"id": "contracts", "label": "合同数据"},
            {"id": "orders", "label": "订单数据"},
            {"id": "invoices", "label": "发票数据"},
            {"id": "bank_summaries", "label": "银行摘要与回款"},
            {"id": "related_transactions", "label": "关联交易摘要"},
        ],
    },
    {
        "id": "financial-check",
        "title": "财务与核验",
        "items": [
            {"id": "latest_metrics", "label": "最新核心财务指标"},
            {"id": "recent_metric_history", "label": "近年财务趋势"},
            {"id": "reconciliation_checks", "label": "勾稽检查结果"},
            {"id": "tax_invoice_checks", "label": "税票一致性核验"},
            {"id": "tax_filings", "label": "纳税申报摘要"},
            {"id": "credit_history", "label": "信用历史摘要"},
        ],
    },
    {
        "id": "external-governance",
        "title": "外部环境与治理",
        "items": [
            {"id": "related_companies", "label": "关联公司关系"},
            {"id": "shareholding_changes", "label": "股权变更记录"},
            {"id": "industry_profile", "label": "行业画像"},
            {"id": "peer_comparisons", "label": "同业比较"},
            {"id": "public_risks", "label": "公开风险事件"},
            {"id": "public_info_enrichment", "label": "公开信息补全层"},
        ],
    },
    {
        "id": "guarantee-support",
        "title": "担保与补件",
        "items": [
            {"id": "guarantees", "label": "担保与抵质押信息"},
            {"id": "due_diligence_materials", "label": "尽调补充材料"},
            {"id": "validation_findings", "label": "核验发现"},
            {"id": "recommendation", "label": "授信建议与缓释要求"},
        ],
    },
]
REVIEW_JUDGEMENT_GUIDANCE = [
    {
        "id": "subject-authenticity",
        "severity": "high",
        "source_type": "主体准入",
        "judgement_focus": "需判断借款主体资格、工商状态和基础证照是否真实有效。",
        "required_data": "营业执照、统一社会信用代码、法定代表人身份证明、公司章程、最新工商信息。",
        "improvement_direction": "补充主体证照与工商核验截图，在报告中写明核验日期与结果。",
    },
    {
        "id": "control-governance",
        "severity": "medium",
        "source_type": "治理结构",
        "judgement_focus": "需判断实际控制关系、股权结构及治理稳定性是否清晰可识别。",
        "required_data": "股权结构图、股权变更记录、董事监事高管名单、关联方清单。",
        "improvement_direction": "补充实控链路说明及治理变化影响，避免只罗列名称不形成判断。",
    },
    {
        "id": "business-substance",
        "severity": "high",
        "source_type": "经营真实性",
        "judgement_focus": "需判断主营业务是否真实持续，经营闭环是否能被合同、订单、发票和回款印证。",
        "required_data": "销售合同、订单、发票、银行流水、主要客户清单与回款记录。",
        "improvement_direction": "将合同、订单、发票与回款链路串起来，明确真实性支撑是否充分。",
    },
    {
        "id": "repayment-source",
        "severity": "high",
        "source_type": "还款来源",
        "judgement_focus": "需判断第一还款来源是否稳定、可验证，是否依赖单一客户或非经常性资金。",
        "required_data": "主营收入结构、核心客户回款、经营现金流、未来订单储备。",
        "improvement_direction": "补充第一还款来源拆解及验证依据，避免只写概括性表述。",
    },
    {
        "id": "cashflow-quality",
        "severity": "high",
        "source_type": "财务质量",
        "judgement_focus": "需判断利润与经营现金流是否匹配，现金回收质量是否支持授信承接。",
        "required_data": "近年利润表、现金流量表、银行流水摘要、应收账款账龄。",
        "improvement_direction": "补写利润到现金流转化逻辑，并解释偏离原因。",
    },
    {
        "id": "leverage-pressure",
        "severity": "medium",
        "source_type": "偿债能力",
        "judgement_focus": "需判断现有债务、杠杆和短债压力是否可控。",
        "required_data": "总负债结构、短期债务、征信余额、到期债务安排、外部融资情况。",
        "improvement_direction": "增加债务期限结构和流动性安排分析，明确短期偿债压力。",
    },
    {
        "id": "tax-invoice-consistency",
        "severity": "medium",
        "source_type": "税票核验",
        "judgement_focus": "需判断收入、开票和银行回款之间是否基本一致，是否存在明显偏差。",
        "required_data": "纳税申报、发票汇总、银行回款、收入确认口径说明。",
        "improvement_direction": "对差异比例较高的月份或客户补充解释，并落到报告核验章节。",
    },
    {
        "id": "industry-policy",
        "severity": "medium",
        "source_type": "行业外部环境",
        "judgement_focus": "需判断行业景气度、监管政策和竞争格局对客户经营的影响。",
        "required_data": "行业研究、政策文件、公开风险、同业比较、企业公开年报摘要。",
        "improvement_direction": "把政策、周期和竞争影响写成授信语言，而不是泛泛行业描述。",
    },
    {
        "id": "related-party-risk",
        "severity": "medium",
        "source_type": "关联交易",
        "judgement_focus": "需判断关联方交易、资金往来和控制链路是否会影响资金独立性。",
        "required_data": "关联方名单、关联交易摘要、往来余额、历史资金拆借说明。",
        "improvement_direction": "补充关联交易定价、公允性和是否存在资金占用风险的判断。",
    },
    {
        "id": "public-risk-event",
        "severity": "medium",
        "source_type": "公开风险",
        "judgement_focus": "需判断处罚、涉诉、舆情或重大公开事件是否影响授信结论。",
        "required_data": "处罚信息、涉诉信息、舆情摘要、年报重大事项、管理层变动。",
        "improvement_direction": "把公开事件与信用影响、期限安排和缓释措施联动起来写。",
    },
    {
        "id": "guarantee-validity",
        "severity": "high",
        "source_type": "担保缓释",
        "judgement_focus": "需判断保证、抵押或质押安排是否真实、有效、足值并可执行。",
        "required_data": "担保承诺、抵质押权属材料、评估报告、押品状态说明。",
        "improvement_direction": "补充权属、估值和处置可行性说明，不要只写押品名称。",
    },
    {
        "id": "compliance-boundary",
        "severity": "medium",
        "source_type": "内外规符合性",
        "judgement_focus": "需判断业务是否符合外规、内部制度和准入口径要求。",
        "required_data": "外规文件、银行内规、准入政策、行业限制清单、产品适配口径。",
        "improvement_direction": "在报告结论里明确列出符合项、待补项和需人工审批判断项。",
    },
    {
        "id": "use-of-proceeds",
        "severity": "high",
        "source_type": "用途合规",
        "judgement_focus": "需判断授信用途是否真实明确、是否与经营规模和交易背景匹配。",
        "required_data": "融资用途说明、用款计划、交易合同、受托支付安排、上下游交易凭证。",
        "improvement_direction": "补写用途真实性核验和支付监测安排，避免只写概括用途。",
    },
    {
        "id": "data-gap",
        "severity": "medium",
        "source_type": "资料完整性",
        "judgement_focus": "需判断当前资料是否足以支持授信结论，是否存在关键缺口。",
        "required_data": "缺失字段清单、未上传原始材料、现场访谈纪要、客户经理补充说明。",
        "improvement_direction": "列出需补充材料清单，并标注这些缺口对结论的影响程度。",
    },
    {
        "id": "final-credit-view",
        "severity": "high",
        "source_type": "综合结论",
        "judgement_focus": "需判断当前尽调结论是否与前文事实、风险点和缓释措施保持一致。",
        "required_data": "尽调正文、风险点列表、缓释措施、授信建议、补件要求。",
        "improvement_direction": "回看全文逻辑闭环，避免前文提示高风险但结论直接放行。",
    },
]
SEEDED_KNOWLEDGE_FILES = [
    {
        "category": "laws",
        "name": "商业银行授信工作尽职指引.pdf",
        "owner": "法律合规岗",
        "description": "授信尽调与审批过程中的合规底线要求。",
        "body": [
            "一、客户经理应核验借款主体、实际控制关系及资金用途真实性。",
            "二、尽职调查报告应明确第一还款来源、第二还款来源及风险缓释安排。",
            "三、对禁入行业、限制行业及重大公开风险客户，应单独说明授信限制条件。",
            "四、引用外部数据、行业材料和公开信息时，应注明来源与使用日期。",
        ],
    },
    {
        "category": "laws",
        "name": "流动资金贷款管理办法.docx",
        "owner": "法律合规岗",
        "description": "流动资金用途、受托支付与贷后要求摘要。",
        "body": [
            "适用范围：适用于借款人日常生产经营周转所需流动资金贷款。",
            "重点要求：贷款用途应真实、明确，与企业经营规模相匹配。",
            "支付要求：大额资金支付应结合受托支付、账户监测和交易凭证留痕管理。",
            "贷后要求：应跟踪资金流向、回款节奏及存量风险变化。",
        ],
    },
    {
        "category": "policies",
        "name": "对公授信准入政策（2026版）.pdf",
        "owner": "风险管理部",
        "description": "行内客户准入、禁入及重点关注行业口径。",
        "body": [
            "准入原则：围绕行业、区域、治理、现金流和风险缓释进行综合准入判断。",
            "禁入口径：对重大违法违规、持续经营异常及高风险名单客户实行禁入。",
            "重点关注：对高杠杆、关联交易复杂、公开处罚频发客户提高审查要求。",
            "材料要求：尽调报告、评级意见、授信方案和外规核验材料应口径一致。",
        ],
    },
    {
        "category": "policies",
        "name": "尽职调查报告写作规范.docx",
        "owner": "授信审批部",
        "description": "尽调报告结构、写法与引用要求。",
        "body": [
            "报告结构：客户概况、经营分析、财务分析、行业环境、抵质押与缓释、风险结论。",
            "写作要求：结论应与数据、材料和核验发现对应，避免仅罗列事实不形成判断。",
            "引用要求：引用知识库、外规、行业研究时应写明文件名称和结论来源。",
            "版本要求：不同版本报告应保留修改说明和新增引用文件清单。",
        ],
    },
    {
        "category": "experience",
        "name": "制造业客户尽调经验要点.pdf",
        "owner": "专家库",
        "description": "制造业客户尽调常见风险点与访谈重点。",
        "body": [
            "关注点一：产销匹配度、核心客户集中度和回款周期是否稳定。",
            "关注点二：应收、存货与现金的联动关系是否合理，是否存在账面利润与现金背离。",
            "关注点三：设备投入、固定资产变动与产能利用率是否匹配。",
            "关注点四：客户经理访谈需补充实际控制人、核心管理层和主要结算账户情况。",
        ],
    },
    {
        "category": "experience",
        "name": "股东关联交易识别案例.docx",
        "owner": "专家库",
        "description": "关联方、资金占用与异常交易识别案例。",
        "body": [
            "案例背景：企业存在多家关联主体，通过往来款和采购销售链路形成资金循环。",
            "识别方法：结合股权结构、银行流水、合同链路和税票信息识别异常闭环。",
            "风险提示：若关联交易定价异常、回款不清晰，应提高授信审查和预审要求。",
            "报告建议：在尽调报告中单列关联方与交易专题，明确是否存在风险传染。",
        ],
    },
]


def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def rows_to_dicts(rows: list[sqlite3.Row]) -> list[dict]:
    return [dict(row) for row in rows]


@lru_cache(maxsize=1)
def get_workbook_index() -> dict:
    wb = load_workbook(XLSX_PATH, read_only=True, data_only=True)
    sheet_names = wb.sheetnames
    company_sheets: dict[str, dict] = defaultdict(lambda: {"sheet_counts": {}, "samples": {}})
    overview_rows: list[dict] = []

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        max_row = ws.max_row if isinstance(ws.max_row, int) and ws.max_row > 0 else 6
        first_rows = list(ws.iter_rows(min_row=1, max_row=min(6, max_row), values_only=True))

        if sheet_name == "Overview":
            for row in ws.iter_rows(min_row=4, values_only=True):
                if row and row[0]:
                    overview_rows.append(
                        {
                            "metric": row[0],
                            "value": row[1],
                            "note": row[3] if len(row) > 3 else None,
                        }
                    )
            continue

        header_row_index = None
        headers = None
        company_code_index = None
        for idx, row in enumerate(first_rows, start=1):
            if row and "company_code" in row:
                header_row_index = idx
                headers = [str(cell) if cell is not None else "" for cell in row]
                company_code_index = headers.index("company_code")
                break

        if header_row_index is None or headers is None or company_code_index is None:
            continue

        for row in ws.iter_rows(min_row=header_row_index + 1, values_only=True):
            if not row or company_code_index >= len(row):
                continue
            company_code = row[company_code_index]
            if not company_code:
                continue
            payload = {
                headers[col_idx]: row[col_idx] if col_idx < len(row) else None
                for col_idx in range(len(headers))
                if headers[col_idx]
            }
            bucket = company_sheets[str(company_code)]
            bucket["sheet_counts"][sheet_name] = bucket["sheet_counts"].get(sheet_name, 0) + 1
            bucket["samples"].setdefault(sheet_name, [])
            if len(bucket["samples"][sheet_name]) < 2:
                bucket["samples"][sheet_name].append(payload)

    return {
        "path": str(XLSX_PATH),
        "sheet_names": sheet_names,
        "overview": overview_rows,
        "companies": company_sheets,
    }


def fetch_one(conn: sqlite3.Connection, query: str, params: tuple = ()) -> dict | None:
    row = conn.execute(query, params).fetchone()
    return dict(row) if row else None


def fetch_all(conn: sqlite3.Connection, query: str, params: tuple = ()) -> list[dict]:
    return rows_to_dicts(conn.execute(query, params).fetchall())


def load_local_config() -> dict:
    if not LOCAL_CONFIG_PATH.exists():
        return {}
    try:
        return json.loads(LOCAL_CONFIG_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash").strip() or "deepseek-v4-flash"


def get_deepseek_api_key() -> str:
    return (os.getenv("DEEPSEEK_API_KEY", "") or load_local_config().get("deepseek_api_key", "")).strip()


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def format_currency(value: float | int | None) -> str:
    if value is None:
        return "—"
    num = float(value)
    if abs(num) >= 100000000:
        return f"{num / 100000000:.2f} 亿元"
    if abs(num) >= 10000:
        return f"{num / 10000:.2f} 万元"
    return f"{num:,.2f} 元"


def compact_text(text: str | None, limit: int = 160) -> str:
    if not text:
        return "暂无补充说明。"
    value = " ".join(str(text).split())
    return value if len(value) <= limit else f"{value[: limit - 1]}…"


def trim_list(items: list[dict] | list | None, limit: int) -> list:
    return list(items or [])[:limit]


def metric_map(metrics: list[dict]) -> dict[str, float | int | None]:
    return {item["metric_code"]: item.get("value") for item in metrics}


def deterministic_rng(company_code: str) -> Random:
    seed = int(hashlib.sha256(company_code.encode("utf-8")).hexdigest()[:16], 16)
    return Random(seed)


def as_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def normalize_ratio(value: object) -> float | None:
    num = as_float(value)
    if num is None:
        return None
    return num / 100 if abs(num) > 1.5 else num


def metric_name_for_code(metric_code: str) -> str:
    mapping = {
        "revenue": "营业收入",
        "gross_profit": "毛利",
        "gross_margin_pct": "毛利率",
        "net_profit": "净利润",
        "cash": "货币资金",
        "operating_cash_flow": "经营现金流",
        "operating_cash_flow_margin_pct": "经营现金流率",
        "accounts_receivable": "应收账款",
        "accounts_payable": "应付账款",
        "inventory": "存货",
        "total_assets": "总资产",
        "total_liabilities": "总负债",
        "net_assets": "净资产",
        "asset_liability_ratio": "资产负债率",
        "capital_expenditure": "资本开支",
        "fixed_assets": "固定资产",
    }
    return mapping.get(metric_code, metric_code)


def metric_unit_for_code(metric_code: str) -> str:
    if metric_code.endswith("_pct") or metric_code == "asset_liability_ratio":
        return "ratio"
    return "CNY"


def upsert_metric(metrics: list[dict], metric_code: str, value: float | None) -> None:
    if value is None:
        return
    for item in metrics:
        if item.get("metric_code") == metric_code:
            item["value"] = value
            item["metric_name"] = item.get("metric_name") or metric_name_for_code(metric_code)
            item["unit"] = item.get("unit") or metric_unit_for_code(metric_code)
            return
    metrics.append(
        {
            "metric_code": metric_code,
            "metric_name": metric_name_for_code(metric_code),
            "metric_category": "poc_generated",
            "value": value,
            "unit": metric_unit_for_code(metric_code),
        }
    )


def normalize_latest_metrics(detail: dict) -> None:
    metrics = list(detail.get("latest_metrics") or [])
    lookup = metric_map(metrics)
    for item in metrics:
        code = item.get("metric_code")
        if code in {"gross_margin_pct", "operating_cash_flow_margin_pct", "asset_liability_ratio"}:
            item["value"] = normalize_ratio(item.get("value"))

    alias_rules = {
        "operating_cashflow": ("operating_cash_flow", as_float),
        "debt_to_assets_pct": ("asset_liability_ratio", normalize_ratio),
        "shareholder_equity": ("net_assets", as_float),
    }
    for source_code, (target_code, transform) in alias_rules.items():
        if lookup.get(target_code) is None and lookup.get(source_code) is not None:
            upsert_metric(metrics, target_code, transform(lookup.get(source_code)))

    lookup = metric_map(metrics)
    revenue = as_float(lookup.get("revenue"))
    gross_profit = as_float(lookup.get("gross_profit"))
    net_profit = as_float(lookup.get("net_profit"))
    operating_cash_flow = as_float(lookup.get("operating_cash_flow"))
    total_assets = as_float(lookup.get("total_assets"))
    total_liabilities = as_float(lookup.get("total_liabilities"))
    net_assets = as_float(lookup.get("net_assets"))
    capital_expenditure = as_float(lookup.get("capital_expenditure"))
    fixed_assets = as_float(lookup.get("fixed_assets"))
    inventory = as_float(lookup.get("inventory"))
    accounts_receivable = as_float(lookup.get("accounts_receivable"))
    accounts_payable = as_float(lookup.get("accounts_payable"))

    raw_metric_map = {item["metric_code"]: as_float(item.get("value")) for item in detail.get("workbook", {}).get("samples", {}).get("financial_metrics", []) if item.get("metric_code")}
    operating_cost = as_float(lookup.get("operating_cost")) or raw_metric_map.get("operating_cost")

    if gross_profit is None and revenue is not None and operating_cost is not None:
        gross_profit = max(revenue - operating_cost, 0)
        upsert_metric(metrics, "gross_profit", gross_profit)
    if lookup.get("gross_margin_pct") is None and revenue and gross_profit is not None:
        upsert_metric(metrics, "gross_margin_pct", gross_profit / revenue)
    if operating_cash_flow is None and raw_metric_map.get("operating_cashflow") is not None:
        operating_cash_flow = raw_metric_map.get("operating_cashflow")
        upsert_metric(metrics, "operating_cash_flow", operating_cash_flow)
    if lookup.get("operating_cash_flow_margin_pct") is None and revenue and operating_cash_flow is not None:
        upsert_metric(metrics, "operating_cash_flow_margin_pct", operating_cash_flow / revenue)
    if lookup.get("asset_liability_ratio") is None and total_assets and total_liabilities is not None:
        upsert_metric(metrics, "asset_liability_ratio", total_liabilities / total_assets)
    if net_assets is None and total_assets is not None and total_liabilities is not None:
        net_assets = total_assets - total_liabilities
        upsert_metric(metrics, "net_assets", net_assets)

    if capital_expenditure is None and revenue is not None:
        industry = detail.get("company", {}).get("industry_category") or ""
        industry_ratio = 0.06 if "软件" in industry or "Software" in industry else 0.11 if "医药" in industry or "医疗" in industry else 0.09
        capital_expenditure = revenue * industry_ratio
        upsert_metric(metrics, "capital_expenditure", capital_expenditure)

    if fixed_assets is None and total_assets is not None:
        industry = detail.get("company", {}).get("industry_category") or ""
        baseline_ratio = 0.08 if "软件" in industry or "信息" in industry or "SaaS" in industry else 0.14 if "医药" in industry or "医疗" in industry else 0.24
        derived_by_assets = total_assets * baseline_ratio
        derived_by_capex = (capital_expenditure or 0) * 1.35
        fixed_assets = max(derived_by_assets, derived_by_capex)
        upsert_metric(metrics, "fixed_assets", fixed_assets)

    if inventory is None and revenue is not None:
        industry = detail.get("company", {}).get("industry_category") or ""
        inventory_ratio = 0.02 if "软件" in industry or "信息" in industry else 0.12 if "医药" in industry else 0.16
        upsert_metric(metrics, "inventory", revenue * inventory_ratio)
    if accounts_receivable is None and revenue is not None:
        industry = detail.get("company", {}).get("industry_category") or ""
        receivable_ratio = 0.08 if "白酒" in industry else 0.14 if "软件" in industry else 0.12
        upsert_metric(metrics, "accounts_receivable", revenue * receivable_ratio)
    if accounts_payable is None and revenue is not None:
        industry = detail.get("company", {}).get("industry_category") or ""
        payable_ratio = 0.07 if "软件" in industry else 0.1 if "制造" in industry or "药" in industry else 0.08
        upsert_metric(metrics, "accounts_payable", revenue * payable_ratio)

    detail["latest_metrics"] = sorted(metrics, key=lambda item: item.get("metric_code", ""))


def is_annual_period_code(period_code: str | None) -> bool:
    if not period_code:
        return False
    if period_code.endswith("FY"):
        return True
    return len(period_code) == 8 and period_code.endswith("1231") and period_code[:4].isdigit()


def normalize_recent_metric_history(detail: dict) -> None:
    items = list(detail.get("recent_metric_history") or [])
    annual_seen: set[tuple[str, str]] = set()
    normalized: list[dict] = []
    for item in items:
        period_code = str(item.get("period_code") or "")
        metric_code = item.get("metric_code")
        if metric_code not in {"revenue", "net_profit", "cash", "accounts_receivable"}:
            continue
        row = dict(item)
        normalized.append(row)
        if is_annual_period_code(period_code):
            annual_seen.add((period_code[:4], metric_code))

    samples = detail.get("workbook", {}).get("samples", {}).get("financial_metrics", [])
    annual_metric_bucket: dict[tuple[str, str], dict] = {}
    for sample in samples:
        period_code = str(sample.get("period_code") or "")
        metric_code = sample.get("metric_code")
        if metric_code in {"revenue", "net_profit", "cash", "accounts_receivable"} and is_annual_period_code(period_code):
            annual_metric_bucket[(period_code[:4], metric_code)] = sample

    for (year, metric_code), sample in annual_metric_bucket.items():
        if (year, metric_code) in annual_seen:
            continue
        normalized.append(
            {
                "period_code": f"{year}1231",
                "metric_code": metric_code,
                "metric_name": sample.get("metric_name") or metric_name_for_code(metric_code),
                "value": sample.get("value"),
                "unit": sample.get("unit") or metric_unit_for_code(metric_code),
            }
        )

    rng = deterministic_rng(detail.get("company", {}).get("company_code") or "poc")
    annual_growth_defaults = {
        "revenue": 0.08,
        "net_profit": 0.06,
        "cash": 0.05,
        "accounts_receivable": 0.07,
    }
    for metric_code, base_growth in annual_growth_defaults.items():
        annual_rows = sorted(
            [item for item in normalized if item.get("metric_code") == metric_code and is_annual_period_code(str(item.get("period_code") or ""))],
            key=lambda item: str(item.get("period_code") or ""),
        )
        if len(annual_rows) >= 2 or not annual_rows:
            continue
        current_row = annual_rows[-1]
        current_year = int(str(current_row.get("period_code"))[:4])
        current_value = as_float(current_row.get("value"))
        if current_value is None:
            continue
        growth = base_growth + (rng.random() - 0.5) * 0.04
        previous_value = current_value / (1 + growth) if current_value >= 0 else current_value * (1 + growth)
        normalized.append(
            {
                "period_code": f"{current_year - 1}1231",
                "metric_code": metric_code,
                "metric_name": current_row.get("metric_name") or metric_name_for_code(metric_code),
                "value": round(previous_value, 2),
                "unit": current_row.get("unit") or metric_unit_for_code(metric_code),
            }
        )

    normalized.sort(key=lambda item: str(item.get("period_code") or ""), reverse=True)
    detail["recent_metric_history"] = normalized[:60]


def supplement_receivables(detail: dict) -> None:
    if detail.get("receivables"):
        return
    metrics = metric_map(detail.get("latest_metrics") or [])
    revenue = as_float(metrics.get("revenue")) or 0
    accounts_receivable = as_float(metrics.get("accounts_receivable")) or revenue * 0.11
    accounts_payable = as_float(metrics.get("accounts_payable")) or revenue * 0.08
    risk_tier = str(detail.get("company", {}).get("risk_tier") or "")
    receivable_days = round(min(max((accounts_receivable / revenue) * 365, 18), 120), 1) if revenue else 48.0
    payable_days = round(min(max((accounts_payable / revenue) * 365, 15), 110), 1) if revenue else 36.0
    top5_ratio = 0.62 if risk_tier == "medium_high" else 0.41 if detail.get("company", {}).get("enterprise_scale") == "listed" else 0.48
    detail["receivables"] = [
        {
            "record_type": "accounts_receivable",
            "balance_cny": accounts_receivable,
            "overdue_over_90d_cny": round(accounts_receivable * (0.06 if risk_tier in {"medium_high", "medium"} else 0.03), 2),
            "top5_ratio_pct": top5_ratio,
            "average_days": receivable_days,
        },
        {
            "record_type": "accounts_payable",
            "balance_cny": accounts_payable,
            "overdue_over_90d_cny": round(accounts_payable * 0.01, 2),
            "top5_ratio_pct": min(top5_ratio - 0.08, 0.58),
            "average_days": payable_days,
        },
    ]


def supplement_bank_summaries(detail: dict) -> None:
    if detail.get("bank_summaries"):
        return
    metrics = metric_map(detail.get("latest_metrics") or [])
    revenue = as_float(metrics.get("revenue")) or 0
    operating_cash_flow = as_float(metrics.get("operating_cash_flow")) or revenue * 0.05
    cash = as_float(metrics.get("cash")) or revenue * 0.08
    inflow = revenue * 0.085
    outflow = inflow - operating_cash_flow / 12
    detail["bank_summaries"] = [
        {
            "inflow_total_cny": round(inflow, 2),
            "outflow_total_cny": round(outflow, 2),
            "net_flow_cny": round(inflow - outflow, 2),
            "average_daily_balance_cny": round(cash * 0.72, 2),
            "top_inflow_counterparty": "核心客户结算账户",
            "top_inflow_amount_cny": round(inflow * 0.31, 2),
            "top_outflow_counterparty": "主要供应商及费用支付",
            "top_outflow_amount_cny": round(outflow * 0.28, 2),
            "transaction_count": 36 if detail.get("company", {}).get("enterprise_scale") == "listed" else 18,
            "large_transaction_count": 5 if detail.get("company", {}).get("enterprise_scale") == "listed" else 3,
            "summary_note": "POC 虚拟补全：根据收入、现金和经营现金流估算月度结算表现。",
        }
    ]


def supplement_tax_checks(detail: dict) -> None:
    if detail.get("tax_invoice_checks"):
        return
    metrics = metric_map(detail.get("latest_metrics") or [])
    revenue = as_float(metrics.get("revenue")) or 0
    if revenue <= 0:
        return
    risk_tier = str(detail.get("company", {}).get("risk_tier") or "")
    declared_ratio = 0.95 if risk_tier in {"medium_high", "medium"} else 0.98
    invoice_ratio = 1.01 if risk_tier == "low" else 1.03
    bank_ratio = 0.92 if risk_tier in {"medium_high", "medium"} else 0.96
    gap_ratio = abs(invoice_ratio - bank_ratio)
    detail["tax_invoice_checks"] = [
        {
            "declared_revenue_cny": round(revenue * declared_ratio, 2),
            "invoiced_amount_cny": round(revenue * invoice_ratio, 2),
            "bank_receipts_cny": round(revenue * bank_ratio, 2),
            "gap_ratio": round(gap_ratio, 4),
            "status": "warn" if gap_ratio >= 0.06 else "pass",
            "check_note": "POC 虚拟补全：以收入与回款节奏推演税票一致性口径。",
        }
    ]


def supplement_reconciliation_checks(detail: dict) -> None:
    if detail.get("reconciliation_checks"):
        return
    metrics = metric_map(detail.get("latest_metrics") or [])
    total_assets = as_float(metrics.get("total_assets")) or 0
    total_liabilities = as_float(metrics.get("total_liabilities")) or 0
    net_assets = as_float(metrics.get("net_assets")) or max(total_assets - total_liabilities, 0)
    revenue = as_float(metrics.get("revenue")) or 0
    gross_profit = as_float(metrics.get("gross_profit")) or revenue * 0.18
    gross_margin_pct = as_float(metrics.get("gross_margin_pct")) or 0.18
    capital_expenditure = as_float(metrics.get("capital_expenditure")) or revenue * 0.08
    fixed_assets = as_float(metrics.get("fixed_assets")) or capital_expenditure * 1.3
    operating_cash_flow = as_float(metrics.get("operating_cash_flow")) or revenue * 0.05
    net_profit = as_float(metrics.get("net_profit")) or revenue * 0.04
    detail["reconciliation_checks"] = [
        {
            "check_name": "资产负债平衡关系",
            "lhs_label": "总资产",
            "lhs_value": total_assets,
            "rhs_label": "总负债+净资产",
            "rhs_value": total_liabilities + net_assets,
            "variance_ratio": 0.0,
            "status": "pass",
            "interpretation": "总资产、负债和权益关系在 POC 补全口径下保持平衡。",
        },
        {
            "check_name": "营业收入减营业成本等于毛利",
            "lhs_label": "收入*毛利率",
            "lhs_value": revenue * gross_margin_pct,
            "rhs_label": "毛利",
            "rhs_value": gross_profit,
            "variance_ratio": 0.0,
            "status": "pass",
            "interpretation": "收入与毛利口径可勾稽，当前未见明显断裂。",
        },
        {
            "check_name": "资本开支与固定资产滚动勾稽",
            "lhs_label": "资本开支",
            "lhs_value": capital_expenditure,
            "rhs_label": "固定资产",
            "rhs_value": fixed_assets,
            "variance_ratio": round(abs(fixed_assets - capital_expenditure) / fixed_assets, 4) if fixed_assets else 0,
            "status": "warn",
            "interpretation": "固定资产与资本开支可形成解释链，但仍建议结合产能投放节奏补充说明。",
        },
        {
            "check_name": "净利润向经营现金流转化",
            "lhs_label": "净利润",
            "lhs_value": net_profit,
            "rhs_label": "经营现金流",
            "rhs_value": operating_cash_flow,
            "variance_ratio": round(abs(operating_cash_flow - net_profit) / abs(net_profit), 4) if net_profit else 0,
            "status": "warn" if operating_cash_flow < net_profit else "pass",
            "interpretation": "利润与现金流需要结合回款、备货和付款节奏一并解释。",
        },
    ]


def supplement_public_risks(detail: dict) -> None:
    if detail.get("public_risks"):
        return
    company = detail.get("company", {})
    subindustry = company.get("subindustry") or company.get("industry_category") or "所属行业"
    company_name = company.get("name") or "目标企业"
    detail["public_risks"] = [
        {
            "event_type": "regulatory_attention",
            "severity": "medium" if company.get("risk_tier") != "low" else "low",
            "event_date": "2025-11-18",
            "title": f"{subindustry}景气波动与价格竞争持续受关注",
            "summary": f"POC 虚拟补全：结合 {company_name} 所处行业生成的外部环境观察项。",
            "public_source_name": "公开行业跟踪",
        },
        {
            "event_type": "information_disclosure",
            "severity": "low",
            "event_date": "2025-08-26",
            "title": "公开信息提示持续关注应收、库存或费用投放变化",
            "summary": "POC 虚拟补全：用于演示外部风险页在公开资料模式下的完整展示。",
            "public_source_name": "公开信息整理",
        },
    ]


def supplement_related_data(detail: dict) -> None:
    company = detail.get("company", {})
    company_code = company.get("company_code")
    company_name = company.get("name") or "目标企业"
    company_short = company_name.replace("股份有限公司", "").replace("有限公司", "")
    risk_tier = str(company.get("risk_tier") or "")
    if company_code == "002594.SZ" and not detail.get("related_companies"):
        detail["related_companies"] = [
            {
                "related_company_name": "融捷投资控股集团有限公司",
                "relation_type": "shareholder",
                "risk_flag": 1,
                "control_path": "公开股东及关联方口径",
                "note": "POC 映射：作为重要股东关联主体展示股东关系与潜在关联交易观察维度。",
            },
            {
                "related_company_name": "弗迪电池有限公司",
                "relation_type": "core_subsidiary",
                "risk_flag": 0,
                "control_path": "公开业务板块及子公司口径",
                "note": "POC 映射：作为核心动力电池平台，用于展示产业链内业务协同关系。",
            },
            {
                "related_company_name": "比亚迪汽车工业有限公司",
                "relation_type": "major_operating_entity",
                "risk_flag": 0,
                "control_path": "公开业务运营平台口径",
                "note": "POC 映射：作为整车制造核心经营平台，用于展示经营主体与授信支持关系。",
            },
        ]
    elif not detail.get("related_companies"):
        detail["related_companies"] = [
            {
                "related_company_name": f"{company_short}供应链管理有限公司",
                "relation_type": "affiliate_supplier",
                "risk_flag": 1 if risk_tier in {"medium", "medium_high"} else 0,
                "control_path": "POC 虚拟补全",
                "note": "POC 虚拟补全：作为长期合作的供应链/渠道协同主体，用于展示关联关系与交易链路。",
            }
        ]
    if company_code == "002594.SZ":
        revenue = as_float(metric_map(detail.get("latest_metrics") or []).get("revenue")) or 0
        desired_transactions = [
            {
                "related_party_name": "融捷投资控股集团有限公司",
                "relation_type": "shareholder",
                "transaction_type": "equity_relation",
                "transaction_amount_cny": round(revenue * 0.0008, 2),
                "revenue_or_cost_ratio": 0.0008,
                "pricing_comment": "POC 映射：股东关系项，重点用于提示是否存在关联资金往来及信息披露义务。",
                "risk_level": "medium",
            },
            {
                "related_party_name": "弗迪电池有限公司",
                "relation_type": "core_subsidiary",
                "transaction_type": "component_supply",
                "transaction_amount_cny": round(revenue * 0.036, 2),
                "revenue_or_cost_ratio": 0.036,
                "pricing_comment": "POC 映射：作为电池业务核心平台，展示集团内部业务协同及定价合理性核验需求。",
                "risk_level": "low",
            },
            {
                "related_party_name": "比亚迪汽车工业有限公司",
                "relation_type": "major_operating_entity",
                "transaction_type": "vehicle_production_settlement",
                "transaction_amount_cny": round(revenue * 0.052, 2),
                "revenue_or_cost_ratio": 0.052,
                "pricing_comment": "POC 映射：作为核心经营平台，展示内部结算与业务链路一致性核验场景。",
                "risk_level": "low",
            },
        ]
        existing_names = {item.get("related_party_name") for item in detail.get("related_transactions") or []}
        additions = [item for item in desired_transactions if item["related_party_name"] not in existing_names]
        if not detail.get("related_transactions"):
            detail["related_transactions"] = desired_transactions
        elif additions:
            detail["related_transactions"] = (detail.get("related_transactions") or []) + additions
    elif not detail.get("related_transactions"):
        detail["related_transactions"] = [
            {
                "related_party_name": detail["related_companies"][0]["related_company_name"],
                "relation_type": detail["related_companies"][0]["relation_type"],
                "transaction_type": "purchase",
                "transaction_amount_cny": round((as_float(metric_map(detail.get("latest_metrics") or []).get("revenue")) or 0) * 0.08, 2),
                "revenue_or_cost_ratio": 0.08,
                "pricing_comment": "POC 虚拟补全：定价整体处于可解释区间，结算节奏需结合回款链路说明。",
                "risk_level": "medium" if risk_tier in {"medium", "medium_high"} else "low",
            }
        ]
    if not detail.get("shareholding_changes"):
        detail["shareholding_changes"] = [
            {
                "change_date": "2025-06-30",
                "change_type": "equity_transfer",
                "declared_reason": "优化长期激励和产业协同安排",
                "risk_comment": "POC 虚拟补全：股权结构有小幅调整，但控制权总体保持稳定。",
            }
        ]


def supplement_industry_profile(detail: dict) -> None:
    if detail.get("industry_profile"):
        return
    company = detail.get("company", {})
    subindustry = company.get("subindustry") or company.get("industry_category") or "所属行业"
    lifecycle = "成熟期" if company.get("enterprise_scale") == "listed" else "成长期"
    detail["industry_profile"] = {
        "subindustry": subindustry,
        "industry_category": company.get("industry_category") or subindustry,
        "policy_direction": f"POC 虚拟补全：{subindustry}整体受产业升级、合规治理和竞争格局变化影响。",
        "lifecycle_stage": lifecycle,
        "benchmark_note": "POC 虚拟补全：建议结合政策、需求、价格竞争和客户集中度综合判断行业风险。",
        "update_cycle": "季度更新",
    }


def supplement_case_and_guarantee(detail: dict) -> None:
    company = detail.get("company", {})
    company_code = company.get("company_code")
    case_info = detail.get("case_info") or {}
    if not case_info.get("product_type"):
        case_info["product_type"] = "public_pack_analysis" if company.get("enterprise_scale") == "listed" else "micro_enterprise_credit"
    if not case_info.get("primary_repayment_source") or case_info.get("primary_repayment_source") == "Not applicable":
        case_info["primary_repayment_source"] = "主营业务经营性现金流与核心客户回款"
    if not case_info.get("guarantee_mode") or case_info.get("guarantee_mode") in {"none", "Not applicable"}:
        case_info["guarantee_mode"] = "核心主体信用+持续信息披露跟踪"
    detail["case_info"] = case_info

    guarantor_names = {item.get("full_name") for item in detail.get("people") or [] if item.get("is_guarantor")}
    if company_code == "002594.SZ":
        additions = []
        if "比亚迪汽车工业有限公司" not in guarantor_names:
            additions.append(
                {
                    "full_name": "比亚迪汽车工业有限公司",
                    "role_type": "guarantor",
                    "title": "核心经营平台 / POC 映射授信支持主体",
                    "equity_ratio": None,
                    "voting_ratio": None,
                    "is_actual_controller": 0,
                    "is_guarantor": 1,
                    "education_level": None,
                    "birth_year": None,
                }
            )
        if "弗迪电池有限公司" not in guarantor_names:
            additions.append(
                {
                    "full_name": "弗迪电池有限公司",
                    "role_type": "guarantor",
                    "title": "核心产业子公司 / POC 映射增信协同主体",
                    "equity_ratio": None,
                    "voting_ratio": None,
                    "is_actual_controller": 0,
                    "is_guarantor": 1,
                    "education_level": None,
                    "birth_year": None,
                }
            )
        if additions:
            detail.setdefault("people", []).extend(additions)
    elif not guarantor_names:
        detail.setdefault("people", []).append(
            {
                "full_name": "授信支持安排",
                "role_type": "guarantor",
                "title": "增信支持方",
                "equity_ratio": None,
                "voting_ratio": None,
                "is_actual_controller": 0,
                "is_guarantor": 1,
                "education_level": None,
                "birth_year": None,
            }
        )

    if company_code == "002594.SZ" and not detail.get("guarantees"):
        metrics = metric_map(detail.get("latest_metrics") or [])
        total_assets = as_float(metrics.get("total_assets")) or 0
        accounts_receivable = as_float(metrics.get("accounts_receivable")) or 0
        detail["guarantees"] = [
            {
                "guarantee_type": "credit_support",
                "asset_name": "主体信用及持续监测安排",
                "asset_category": "信用增信",
                "appraised_value_cny": round(total_assets * 0.18, 2) if total_assets else 0,
                "pledge_rate": 0.0,
                "lien_status": "clear",
                "guarantee_status": "active",
                "notes": "POC 映射：参考公开主体规模与信用资质，作为主信用增信展示。",
            },
            {
                "guarantee_type": "receivable_pool",
                "asset_name": "核心应收账款质押池（POC 映射）",
                "asset_category": "应收账款质押",
                "appraised_value_cny": round(accounts_receivable * 0.72, 2) if accounts_receivable else 0,
                "pledge_rate": 0.46,
                "lien_status": "clear",
                "guarantee_status": "active",
                "notes": "POC 映射：用于展示整车及电池业务回款形成的应收账款缓释安排。",
            },
            {
                "guarantee_type": "subsidiary_support",
                "asset_name": "核心子公司经营支持承诺（POC 映射）",
                "asset_category": "经营支持函",
                "appraised_value_cny": round(total_assets * 0.11, 2) if total_assets else 0,
                "pledge_rate": 0.0,
                "lien_status": "clear",
                "guarantee_status": "active",
                "notes": "POC 映射：结合公开子公司板块地位，模拟经营支持型增信安排。",
            },
        ]
    elif not detail.get("guarantees"):
        metrics = metric_map(detail.get("latest_metrics") or [])
        total_assets = as_float(metrics.get("total_assets")) or 0
        detail["guarantees"] = [
            {
                "guarantee_type": "credit_support",
                "asset_name": "主体信用及持续监测安排",
                "asset_category": "信用增信",
                "appraised_value_cny": round(total_assets * 0.18, 2) if total_assets else 0,
                "pledge_rate": 0.0,
                "lien_status": "clear",
                "guarantee_status": "active",
                "notes": "POC 虚拟补全：用于演示授信缓释、持续监测与补充增信安排。",
            }
        ]


def supplement_credit_history(detail: dict) -> None:
    company = detail.get("company", {})
    company_code = company.get("company_code")
    if company_code == "002594.SZ":
        existing = detail.get("credit_history") or []
        existing_subjects = {item.get("subject_name") for item in existing}
        additions = []
        if "比亚迪汽车工业有限公司" not in existing_subjects:
            additions.append(
                {
                    "subject_type": "enterprise",
                    "subject_name": "比亚迪汽车工业有限公司",
                    "credit_channel": "public_operating_platform_profile",
                    "account_count": 8,
                    "outstanding_balance_cny": 15264000000.0,
                    "overdue_count": 0,
                    "max_overdue_bucket": "none",
                    "hard_inquiry_3m": 0,
                    "hard_inquiry_6m": 1,
                    "credit_assessment": "stable",
                    "summary_note": "POC 映射：核心经营平台公开口径下信用记录整体稳定，用于展示担保支持方信用摘要。",
                }
            )
        if "弗迪电池有限公司" not in existing_subjects:
            additions.append(
                {
                    "subject_type": "enterprise",
                    "subject_name": "弗迪电池有限公司",
                    "credit_channel": "public_core_subsidiary_profile",
                    "account_count": 6,
                    "outstanding_balance_cny": 9850000000.0,
                    "overdue_count": 0,
                    "max_overdue_bucket": "none",
                    "hard_inquiry_3m": 0,
                    "hard_inquiry_6m": 0,
                    "credit_assessment": "stable",
                    "summary_note": "POC 映射：核心产业子公司公开口径下信用记录整体可控，用于展示增信协同主体摘要。",
                }
            )
        if additions:
            detail["credit_history"] = existing + additions
            return
    if detail.get("credit_history"):
        return
    metrics = metric_map(detail.get("latest_metrics") or [])
    total_liabilities = as_float(metrics.get("total_liabilities")) or 0
    detail["credit_history"] = [
        {
            "subject_type": "company",
            "subject_name": detail.get("company", {}).get("name") or "目标企业",
            "credit_channel": "银行授信及公开融资",
            "account_count": 4 if detail.get("company", {}).get("enterprise_scale") == "listed" else 2,
            "outstanding_balance_cny": round(total_liabilities * 0.42, 2) if total_liabilities else 0,
            "overdue_count": 0,
            "max_overdue_bucket": "无",
            "hard_inquiry_3m": 1,
            "hard_inquiry_6m": 2,
            "credit_assessment": "正常",
            "summary_note": "POC 虚拟补全：当前未见明显逾期，但建议持续跟踪存量债务与融资节奏。",
        }
    ]


def supplement_peer_comparisons(detail: dict) -> None:
    if detail.get("peer_comparisons"):
        return
    metrics = metric_map(detail.get("latest_metrics") or [])
    detail["peer_comparisons"] = [
        {
            "company_value": as_float(metrics.get("gross_margin_pct")) or 0.18,
            "variance_pct": -0.02,
            "percentile_bucket": "行业中位附近",
            "narrative": "POC 虚拟补全：盈利能力接近行业中位，需进一步关注现金转化和营运资本压力。",
        },
        {
            "company_value": as_float(metrics.get("asset_liability_ratio")) or 0.55,
            "variance_pct": 0.03,
            "percentile_bucket": "略高于行业中位",
            "narrative": "POC 虚拟补全：杠杆水平略高于同业中位，建议审慎判断持续承压能力。",
        },
    ]


def supplement_report_sections(detail: dict) -> None:
    if detail.get("report_sections"):
        return
    company_name = detail.get("company", {}).get("name") or "目标企业"
    detail["report_sections"] = [
        {"section_code": "company_profile", "section_title": "企业概况", "content": f"{company_name} 已完成 POC 虚拟资料补全，可用于页面演示与授信判断结构展示。", "confidence": 0.8},
        {"section_code": "operations", "section_title": "经营分析", "content": "当前经营、回款与外部环境已补充为可展示口径，便于前端呈现专题信息。", "confidence": 0.78},
        {"section_code": "finance", "section_title": "财务分析", "content": "财务页已按结论先行和多维分析方式展示，缺失指标由 POC 补全器推演。", "confidence": 0.76},
    ]


def supplement_poc_detail(detail: dict) -> dict:
    normalize_latest_metrics(detail)
    normalize_recent_metric_history(detail)
    supplement_industry_profile(detail)
    supplement_case_and_guarantee(detail)
    supplement_receivables(detail)
    supplement_bank_summaries(detail)
    supplement_tax_checks(detail)
    supplement_reconciliation_checks(detail)
    supplement_public_risks(detail)
    supplement_related_data(detail)
    supplement_credit_history(detail)
    supplement_peer_comparisons(detail)
    supplement_report_sections(detail)
    return detail


def ensure_knowledge_library_dirs() -> None:
    KNOWLEDGE_FILES_DIR.mkdir(parents=True, exist_ok=True)


def save_knowledge_manifest(files: list[dict], next_id: int) -> None:
    ensure_knowledge_library_dirs()
    KNOWLEDGE_META_PATH.write_text(
        json.dumps({"next_id": next_id, "files": files}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def create_demo_pdf(path: Path, title: str, body_lines: list[str]) -> None:
    pdf_bytes = build_pdf_bytes(title, "知识库样例", "\n".join(body_lines))
    path.write_bytes(pdf_bytes)


def create_demo_docx(path: Path, title: str, body_lines: list[str]) -> None:
    document = Document()
    document.add_heading(title, 0)
    for line in body_lines:
      document.add_paragraph(line)
    document.save(path)


def create_seeded_knowledge_files() -> tuple[list[dict], int]:
    ensure_knowledge_library_dirs()
    files: list[dict] = []
    next_id = 1
    for item in SEEDED_KNOWLEDGE_FILES:
        file_id = f"kb-{next_id}"
        next_id += 1
        suffix = item["name"].split(".")[-1].lower()
        stored_name = f"{file_id}.{suffix}"
        file_path = KNOWLEDGE_FILES_DIR / stored_name
        if suffix == "pdf":
            create_demo_pdf(file_path, item["name"], item["body"])
        else:
            create_demo_docx(file_path, item["name"], item["body"])
        content_text = "\n".join(item["body"])
        files.append(
            {
                "id": file_id,
                "name": item["name"],
                "category": item["category"],
                "uploaded_at": now_iso(),
                "owner": item["owner"],
                "description": item["description"],
                "format": suffix,
                "stored_name": stored_name,
                "content_text": content_text,
            }
        )
    save_knowledge_manifest(files, next_id)
    return files, next_id


def load_public_enrichment(company_code: str) -> dict:
    base = PROJECT_ROOT / "artifacts" / "company_files" / company_code
    files = {
        "enterprise_profile": "enterprise_profile.md",
        "annual_reports": "annual_reports.md",
        "industry_insights": "industry_insights.md",
        "public_risks": "public_risks.md",
        "governance_summary": "governance_summary.md",
        "report_summary": "report_summary.md",
        "annual_financial_overview": "annual_financial_overview.md",
    }
    payload: dict[str, str] = {}
    for key, filename in files.items():
        path = base / filename
        payload[key] = compact_text(path.read_text(encoding="utf-8"), 2200) if path.exists() else ""
    payload["available"] = any(bool(value.strip()) for value in payload.values())
    return payload


def get_skill_pipeline_binding(company_code: str) -> dict | None:
    if not all(get_skill_dir(skill_name).exists() for skill_name in GLOBAL_SKILL_PIPELINE["skills"]):
        return None
    mock_binding = SKILL_PIPELINE_MOCK_BINDINGS.get(company_code) or {}
    binding = {
        "company_code": company_code,
        "label": GLOBAL_SKILL_PIPELINE["label"],
        "skills": list(GLOBAL_SKILL_PIPELINE["skills"]),
        "entry_mode": GLOBAL_SKILL_PIPELINE["entry_mode"],
        "mock_vdr_dir": str(mock_binding["mock_vdr_dir"]) if mock_binding.get("mock_vdr_dir") else None,
        "company_name": mock_binding.get("company_name"),
    }
    if mock_binding:
        binding["label"] = "通用双 Skill 报告生成（灵犀含 Mock VDR 兜底）"
    return binding


def has_antigravity_mock_pack(company_code: str) -> bool:
    mock_binding = SKILL_PIPELINE_MOCK_BINDINGS.get(company_code)
    return bool(mock_binding and mock_binding.get("mock_vdr_dir") and mock_binding["mock_vdr_dir"].exists())


def get_skill_dir(skill_name: str) -> Path:
    return PROJECT_ROOT / ".agent" / "skills" / skill_name


def load_skill_file(skill_name: str, filename: str = "SKILL.md") -> str:
    path = get_skill_dir(skill_name) / filename
    return read_text_file(path)


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def parse_markdown_table(md_text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r"-+", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def parse_money_amount(text: str) -> int | None:
    match = re.search(r"([0-9][0-9,]*)\s*RMB", text)
    if not match:
        return None
    return int(match.group(1).replace(",", ""))


def parse_percentage_amount(text: str) -> float | None:
    match = re.search(r"([0-9]+(?:\.[0-9]+)?)%", text)
    if not match:
        return None
    return float(match.group(1))


def load_antigravity_mock_pack(company_code: str) -> dict:
    binding = SKILL_PIPELINE_MOCK_BINDINGS[company_code]
    base = binding["mock_vdr_dir"]
    business_license_text = read_text_file(base / "business_license.md")
    shareholding_text = read_text_file(base / "shareholding_register.md")
    financial_text = read_text_file(base / "financial_statements.md")
    electricity_text = read_text_file(base / "electricity_bills_2025.txt")
    contract_text = read_text_file(base / "top_customer_contracts.md")
    bank_text = read_text_file(base / "bank_statement_summary.md")
    tax_text = read_text_file(base / "tax_filing_summary.md")
    board_text = read_text_file(base / "board_meeting_minutes_secret.md")
    litigation_text = read_text_file(base / "litigation_and_penalties.md")
    interview_text = read_text_file(base / "management_interview_notes.md")

    table_rows = parse_markdown_table(financial_text)
    income_rows = {
        row[0]: {
            "revenue": int(row[1]),
            "gross_profit": int(row[2]),
            "net_profit": int(row[3]),
            "revenue_growth_pct": float(row[4].replace("%", "")),
        }
        for row in table_rows
        if len(row) >= 5 and row[0] in {"2024", "2025"}
    }

    top_clients = []
    for line in financial_text.splitlines():
        if not line.strip().startswith("- Client"):
            continue
        name_match = re.search(r":\s*(.+?)\s+\(", line)
        revenue_match = re.search(r"贡献收入:\s*([0-9,]+)\s*RMB", line)
        share_match = re.search(r"占比:\s*([0-9.]+)%", line)
        if not name_match:
            continue
        top_clients.append(
            {
                "name": name_match.group(1).strip(),
                "revenue": int(revenue_match.group(1).replace(",", "")) if revenue_match else 0,
                "share_pct": float(share_match.group(1)) if share_match else 0.0,
            }
        )

    assets_match = re.search(r"Total Assets:\s*([0-9,]+)\s*RMB", financial_text)
    liabilities_match = re.search(r"Total Liabilities:\s*([0-9,]+)\s*RMB", financial_text)
    asset_ratio_match = re.search(r"Asset-Liability Ratio:\s*([0-9.]+)%", financial_text)

    month_lines = re.findall(r"([A-Z][a-z]{2}) 2025:\s*([0-9,]+)\s*kWh", electricity_text)
    total_2024_match = re.search(r"Total 2024 Consumption:\s*([0-9,]+)\s*kWh", electricity_text)
    total_2025_match = re.search(r"Total 2025 Consumption:\s*([0-9,]+)\s*kWh", electricity_text)
    power_ratio_match = re.search(r"Calculated Power Fluctuating Ratio:\s*\+?([0-9.]+)%", electricity_text)

    meeting_date_match = re.search(r"会议时间：([0-9]{4}\s*年\s*[0-9]{1,2}\s*月\s*[0-9]{1,2}\s*日)", board_text)
    directors_match = re.search(r"参会董事：(.+)", board_text)
    guarantee_term_match = re.search(r"([0-9]+)\s*年期企业贷款", board_text)
    guarantee_amount_match = re.search(r"申请的\s*([0-9,]+)\s*RMB", board_text)
    related_customer_match = re.search(r"与 “(.+?)” (?:的常年战略|执行年度框架)", board_text)
    guaranteed_company_match = re.search(r"兄弟企业 “(.+?)”", board_text)
    legal_rep_match = re.search(r"法定代表人：(.+)", business_license_text)
    established_match = re.search(r"成立日期：([0-9]{4}年[0-9]{2}月[0-9]{2}日)", business_license_text)
    registered_capital_match = re.search(r"注册资本：([0-9,]+)\s*RMB", business_license_text)
    annual_receipts_match = re.search(r"2025 年可识别客户销售回款合计：([0-9,]+)\s*RMB", bank_text)
    q4_booked_match = re.search(r"Q4 财务确认收入：([0-9,]+)\s*RMB", bank_text)
    q4_receipts_match = re.search(r"Q4 银行实际回款：([0-9,]+)\s*RMB", bank_text)
    tax_declared_match = re.search(r"2025 年累计申报收入：([0-9,]+)\s*RMB", tax_text)

    directors = []
    for raw_item in (directors_match.group(1) if directors_match else "").split("、"):
        item = raw_item.strip()
        if not item:
            continue
        person_match = re.match(r"(.+?)(?:（持股([0-9]+)%）)?$", item)
        if not person_match:
            continue
        directors.append(
            {
                "full_name": person_match.group(1).strip(),
                "equity_ratio": (float(person_match.group(2)) / 100) if person_match.group(2) else 0.0,
            }
        )

    return {
        "base_dir": base,
        "company_name": binding.get("company_name", "目标公司"),
        "business_license_text": business_license_text,
        "shareholding_text": shareholding_text,
        "financial_text": financial_text,
        "electricity_text": electricity_text,
        "contract_text": contract_text,
        "bank_text": bank_text,
        "tax_text": tax_text,
        "board_text": board_text,
        "litigation_text": litigation_text,
        "interview_text": interview_text,
        "income_rows": income_rows,
        "top_clients": top_clients,
        "total_assets": int(assets_match.group(1).replace(",", "")) if assets_match else 0,
        "total_liabilities": int(liabilities_match.group(1).replace(",", "")) if liabilities_match else 0,
        "asset_liability_ratio_pct": float(asset_ratio_match.group(1)) if asset_ratio_match else 0.0,
        "monthly_power_kwh": [{"month": month, "value": int(value.replace(",", ""))} for month, value in month_lines],
        "power_2024_kwh": int(total_2024_match.group(1).replace(",", "")) if total_2024_match else 0,
        "power_2025_kwh": int(total_2025_match.group(1).replace(",", "")) if total_2025_match else 0,
        "power_growth_pct": float(power_ratio_match.group(1)) if power_ratio_match else 0.0,
        "meeting_date": meeting_date_match.group(1).replace(" ", "") if meeting_date_match else "2025年10月15日",
        "directors": directors,
        "director_name": directors[0]["full_name"] if directors else "实际控制人",
        "related_customer_name": related_customer_match.group(1).strip() if related_customer_match else (top_clients[0]["name"] if top_clients else "核心客户"),
        "guaranteed_company_name": guaranteed_company_match.group(1).strip() if guaranteed_company_match else "关联兄弟企业",
        "guarantee_amount_rmb": int(guarantee_amount_match.group(1).replace(",", "")) if guarantee_amount_match else 0,
        "guarantee_term_years": int(guarantee_term_match.group(1)) if guarantee_term_match else 3,
        "legal_representative": legal_rep_match.group(1).strip() if legal_rep_match else "",
        "established_on_text": established_match.group(1) if established_match else "",
        "registered_capital_rmb": int(registered_capital_match.group(1).replace(",", "")) if registered_capital_match else 0,
        "annual_receipts_rmb": int(annual_receipts_match.group(1).replace(",", "")) if annual_receipts_match else 0,
        "q4_booked_revenue_rmb": int(q4_booked_match.group(1).replace(",", "")) if q4_booked_match else 0,
        "q4_receipts_rmb": int(q4_receipts_match.group(1).replace(",", "")) if q4_receipts_match else 0,
        "tax_declared_revenue_rmb": int(tax_declared_match.group(1).replace(",", "")) if tax_declared_match else 0,
    }


def build_antigravity_audit_summary(company_code: str) -> dict:
    pack = load_antigravity_mock_pack(company_code)
    revenue_2024 = pack["income_rows"].get("2024", {}).get("revenue", 0)
    revenue_2025 = pack["income_rows"].get("2025", {}).get("revenue", 0)
    revenue_growth_pct = ((revenue_2025 - revenue_2024) * 100 / revenue_2024) if revenue_2024 else 0.0
    top_customer = pack["top_clients"][0] if pack["top_clients"] else {"name": "", "revenue": 0, "share_pct": 0.0}
    source_prefix = f".agent/mock_vdr/{pack['base_dir'].name}"

    findings = []
    if revenue_growth_pct > 15 and pack["power_growth_pct"] < 3:
        findings.append(
            {
                "severity": "高",
                "finding_title": "主营业务收入真实性存疑",
                "problem_definition": "收入增幅明显高于用电量增幅，经营增长与物理产能代理指标未同步。",
                "financial_impact": f"新增收入 {revenue_2025 - revenue_2024} 元仍需补充产能支持、外协支撑或收入确认时点说明。",
                "evidence_chain": [
                    f"2024 年收入为 {revenue_2024} 元，2025 年收入为 {revenue_2025} 元。",
                    f"2024 年总用电量为 {pack['power_2024_kwh']} 千瓦时，2025 年总用电量为 {pack['power_2025_kwh']} 千瓦时。",
                ],
                "source_refs": [
                    {"file": f"{source_prefix}/financial_statements.md", "line_hint": "Income Statement Highlights"},
                    {"file": f"{source_prefix}/electricity_bills_2025.txt", "line_hint": "Calculated Power Fluctuating Ratio"},
                ],
            }
        )

    if (
        top_customer["name"]
        and top_customer["name"] in pack["board_text"]
        and (
            top_customer["name"] in pack["shareholding_text"]
            or "家庭关联方" in pack["board_text"]
            or "家庭控制" in pack["board_text"]
        )
        and "No related-party transactions requiring special disclosure" in pack["financial_text"]
    ):
        findings.append(
            {
                "severity": "高",
                "finding_title": "未披露重大关联交易风险",
                "problem_definition": "第一大客户与实控人家庭存在控制重合，但对外口径未充分披露关联关系。",
                "financial_impact": f"至少 {top_customer['revenue']} 元收入涉及关联交易公允性、客户独立性和定价合理性复核。",
                "evidence_chain": [
                    f"{top_customer['name']} 贡献收入 {top_customer['revenue']} 元，占 2025 年收入的 {top_customer['share_pct']:.1f}%。",
                    f"内部材料显示 {top_customer['name']} 由 {pack['director_name']} 家庭控制。",
                ],
                "source_refs": [
                    {"file": f"{source_prefix}/financial_statements.md", "line_hint": "Top Revenue Contributors (2025)"},
                    {"file": f"{source_prefix}/shareholding_register.md", "line_hint": "实控人及家庭对外控制企业"},
                    {"file": f"{source_prefix}/board_meeting_minutes_secret.md", "line_hint": "议题一：关于核心客户关系说明"},
                ],
            }
        )

    if (
        pack["guarantee_amount_rmb"]
        and "连带责任担保" in pack["board_text"]
        and ("暂不在对外资产负债表" in pack["board_text"] or "不对外披露" in pack["board_text"])
    ):
        findings.append(
            {
                "severity": "高",
                "finding_title": "未披露表外连带责任风险",
                "problem_definition": "公司已批准对关联兄弟企业提供全额连带责任担保，但要求财务不对外披露。",
                "financial_impact": f"潜在或有负债敞口为 {pack['guarantee_amount_rmb']} 元，后续可能引发代偿、资产冻结及诉讼风险。",
                "evidence_chain": [
                    f"董事会批准为关联企业提供 {pack['guarantee_amount_rmb']} 元、期限 {pack['guarantee_term_years']} 年的全额连带责任担保。",
                    "被担保企业已出现民间借贷利息逾期。",
                    "会议纪要明确要求财务暂不对外披露该或有负债。",
                ],
                "source_refs": [
                    {"file": f"{source_prefix}/board_meeting_minutes_secret.md", "line_hint": "议题二：关于兄弟企业授信担保安排"},
                    {"file": f"{source_prefix}/board_meeting_minutes_secret.md", "line_hint": "会议记录补充"},
                ],
            }
        )

    if pack["annual_receipts_rmb"] and pack["tax_declared_revenue_rmb"]:
        findings.append(
            {
                "severity": "中",
                "finding_title": "收入回款税票勾稽差异",
                "problem_definition": "财务报表收入高于可识别银行回款及税务申报收入。",
                "financial_impact": (
                    f"2025 年财务收入较可识别回款高 {revenue_2025 - pack['annual_receipts_rmb']} 元，"
                    f"较税务申报收入高 {revenue_2025 - pack['tax_declared_revenue_rmb']} 元。"
                ),
                "evidence_chain": [
                    f"2025 年可识别银行回款为 {pack['annual_receipts_rmb']} 元。",
                    f"2025 年税务申报收入为 {pack['tax_declared_revenue_rmb']} 元。",
                ],
                "source_refs": [
                    {"file": f"{source_prefix}/bank_statement_summary.md", "line_hint": "全年回款口径汇总"},
                    {"file": f"{source_prefix}/tax_filing_summary.md", "line_hint": "增值税申报口径"},
                ],
            }
        )

    return {
        "company_name": pack["company_name"],
        "analysis_scope": source_prefix,
        "computed_metrics": {
            "revenue_2024_rmb": revenue_2024,
            "revenue_2025_rmb": revenue_2025,
            "revenue_growth_pct": round(revenue_growth_pct, 1),
            "power_consumption_2024_kwh": pack["power_2024_kwh"],
            "power_consumption_2025_kwh": pack["power_2025_kwh"],
            "power_usage_growth_pct": round(pack["power_growth_pct"], 1),
            "largest_customer_name": top_customer["name"],
            "largest_customer_revenue_rmb": top_customer["revenue"],
            "largest_customer_revenue_share_pct": top_customer["share_pct"],
            "undisclosed_guarantee_amount_rmb": pack["guarantee_amount_rmb"],
            "bank_receipts_2025_rmb": pack["annual_receipts_rmb"],
            "tax_filed_revenue_2025_rmb": pack["tax_declared_revenue_rmb"],
        },
        "overall_risk_level": "高" if any(item["severity"] == "高" for item in findings) else "中",
        "findings": findings,
        "writer_handoff": {
            "report_title": "尽职调查报告",
            "must_highlight_first": [item["finding_title"] for item in findings if item["severity"] == "高"],
        },
    }


def normalize_skill_severity(value: object) -> str:
    text = str(value or "").strip().lower()
    if text in {"高", "high", "red", "critical"}:
        return "高"
    if text in {"低", "low", "green"}:
        return "低"
    return "中"


def build_structured_skill_audit_summary(company_code: str, detail: dict | None = None) -> dict:
    detail = detail or get_company_detail_base_payload(company_code) or {}
    company = detail.get("company") or {}
    findings: list[dict] = []

    for item in detail.get("validation_findings") or []:
        title = item.get("finding_title") or item.get("title") or item.get("rule_name") or "结构化核验发现"
        severity = normalize_skill_severity(item.get("severity") or item.get("risk_level"))
        findings.append(
            {
                "severity": severity,
                "finding_title": title,
                "problem_definition": item.get("finding_summary") or item.get("summary") or item.get("description") or "结构化数据存在待复核事项。",
                "financial_impact": item.get("impact") or item.get("risk_impact") or "可能影响授信风险判断，需结合原始资料进一步核验。",
                "evidence_chain": [
                    value
                    for value in [
                        item.get("evidence"),
                        item.get("source_ref"),
                        item.get("finding_detail"),
                    ]
                    if value
                ],
                "source_refs": [{"file": "structured://validation_findings", "line_hint": title}],
            }
        )

    for item in detail.get("reconciliation_checks") or []:
        if str(item.get("result") or item.get("status") or "").lower() in {"pass", "passed", "ok"}:
            continue
        title = item.get("check_name") or item.get("rule_name") or "财务勾稽差异"
        findings.append(
            {
                "severity": normalize_skill_severity(item.get("severity") or item.get("risk_level")),
                "finding_title": title,
                "problem_definition": item.get("description") or item.get("finding") or "财务、银行流水或税务口径存在勾稽差异。",
                "financial_impact": item.get("impact") or "可能影响收入真实性、现金回收质量和第一还款来源判断。",
                "evidence_chain": [value for value in [item.get("metric_name"), item.get("difference_note"), item.get("source_ref")] if value],
                "source_refs": [{"file": "structured://financial_reconciliation_checks", "line_hint": title}],
            }
        )

    for item in detail.get("tax_invoice_checks") or []:
        if str(item.get("result") or item.get("status") or "").lower() in {"pass", "passed", "ok"}:
            continue
        title = item.get("check_name") or item.get("rule_name") or "税票一致性差异"
        findings.append(
            {
                "severity": normalize_skill_severity(item.get("severity") or item.get("risk_level")),
                "finding_title": title,
                "problem_definition": item.get("description") or "税务申报、发票或收入确认口径存在待解释差异。",
                "financial_impact": item.get("impact") or "可能影响收入确认合规性和纳税申报可信度。",
                "evidence_chain": [value for value in [item.get("difference_note"), item.get("source_ref")] if value],
                "source_refs": [{"file": "structured://tax_invoice_consistency_checks", "line_hint": title}],
            }
        )

    for item in detail.get("public_risks") or []:
        title = item.get("risk_title") or item.get("title") or "公开风险事件"
        findings.append(
            {
                "severity": normalize_skill_severity(item.get("severity") or item.get("risk_level")),
                "finding_title": title,
                "problem_definition": item.get("description") or item.get("summary") or "公开信息中存在待关注风险事件。",
                "financial_impact": item.get("impact") or "需评估其对经营连续性、合规记录和偿债能力的影响。",
                "evidence_chain": [value for value in [item.get("source_name"), item.get("source_ref"), item.get("event_date")] if value],
                "source_refs": [{"file": "structured://public_risks", "line_hint": title}],
            }
        )

    if not findings:
        findings.append(
            {
                "severity": "中",
                "finding_title": "资料完整性与结论一致性待复核",
                "problem_definition": "当前结构化数据未触发明确高风险项，但仍需围绕主体准入、经营真实性、财务质量和担保缓释完成一致性复核。",
                "financial_impact": "若资料缺口未补齐，可能影响授信额度、期限、担保条件及贷后监控安排。",
                "evidence_chain": ["基于当前企业详情、财务指标、核验发现、公开风险和知识库文件形成初步审计输入。"],
                "source_refs": [{"file": "structured://company_detail", "line_hint": "全量企业详情"}],
            }
        )

    high_count = sum(1 for item in findings if item["severity"] == "高")
    return {
        "company_name": company.get("name") or company_code,
        "analysis_scope": "structured_company_payload",
        "computed_metrics": {
            "latest_metric_count": len(detail.get("latest_metrics") or []),
            "validation_finding_count": len(detail.get("validation_findings") or []),
            "public_risk_count": len(detail.get("public_risks") or []),
            "credit_history_count": len(detail.get("credit_history") or []),
        },
        "overall_risk_level": "高" if high_count else "中",
        "findings": findings[:12],
        "writer_handoff": {
            "report_title": "尽职调查报告",
            "must_highlight_first": [item["finding_title"] for item in findings if item["severity"] == "高"],
        },
    }


def build_skill_audit_summary(company_code: str, detail: dict | None = None) -> dict:
    if has_antigravity_mock_pack(company_code):
        return build_antigravity_audit_summary(company_code)
    return build_structured_skill_audit_summary(company_code, detail)


def build_antigravity_report_markdown(company_code: str, version_number: int, detail: dict | None = None) -> str:
    audit = build_antigravity_audit_summary(company_code)
    pack = load_antigravity_mock_pack(company_code)
    metrics = audit["computed_metrics"]
    findings = {item["finding_title"]: item for item in audit["findings"]}
    high_count = sum(1 for item in audit["findings"] if item["severity"] == "高")
    medium_count = sum(1 for item in audit["findings"] if item["severity"] == "中")
    source_prefix = audit["analysis_scope"]
    revenue_issue = findings.get("主营业务收入真实性存疑", {})
    related_issue = findings.get("未披露重大关联交易风险", {})
    liability_issue = findings.get("未披露表外连带责任风险", {})
    cash_tax_issue = findings.get("收入回款税票勾稽差异", {})
    company_name = detail.get("company", {}).get("name") if detail else audit["company_name"]

    return "\n".join(
        [
            "# 尽职调查报告",
            "",
            "## 第一章 基本情况与历史沿革核查",
            "### 风险定性",
            (
                f"{company_name}营业执照显示法定代表人为 {pack['legal_representative']}，成立日期为 {pack['established_on_text']}，注册资本 {pack['registered_capital_rmb']} 元。"
                "根据当前材料，基础工商口径已形成初步画像，但实控人家庭对外控制企业较多，主体独立性仍需结合历史沿革、关联控制链条和治理安排进一步核查"
                f" [{source_prefix}/business_license.md] [{source_prefix}/shareholding_register.md]。"
            ),
            "",
            "### 核查证据",
            "| 项目 | 内容 | 核查说明 |",
            "| --- | --- | --- |",
            f"| 企业名称 | {company_name} | 已同步基础主体信息 |",
            f"| 法定代表人 | {pack['legal_representative']} | 待补充任职与治理稳定性说明 |",
            f"| 成立日期 | {pack['established_on_text']} | 可继续核对历史沿革材料 |",
            f"| 注册资本 | {pack['registered_capital_rmb']} 元 | 需结合章程和工商档案复核 |",
            "",
            "### 财务与合规影响评估",
            "若主体沿革、关联控制链条及治理独立性披露不充分，可能影响授信准入、关联交易识别以及后续责任承担判断。",
            "",
            "### 补充尽调建议",
            "建议补充工商底档、章程、历次股权变更材料及实际控制人说明。",
            "",
            "## 第二章 主营业务与经营模式深挖",
            "### 风险定性",
            (
                f"2025 年收入同比增长 {metrics['revenue_growth_pct']:.1f}%，同期总用电量仅增长 {metrics['power_usage_growth_pct']:.1f}%，"
                "经营增长与物理世界代理指标存在明显背离。若管理层关于委外加工的解释缺少合同、结算单与入库单支持，则主营业务收入真实性仍需进一步核验"
                f" [{source_prefix}/financial_statements.md] [{source_prefix}/electricity_bills_2025.txt] [{source_prefix}/management_interview_notes.md]。"
            ),
            "",
            "### 核查证据",
            "| 核查维度 | 2024 年 | 2025 年 | 变动情况 |",
            "| --- | --- | --- | --- |",
            f"| 营业收入 | {metrics['revenue_2024_rmb']} 元 | {metrics['revenue_2025_rmb']} 元 | {metrics['revenue_growth_pct']:.1f}% |",
            f"| 总用电量 | {metrics['power_consumption_2024_kwh']} 千瓦时 | {metrics['power_consumption_2025_kwh']} 千瓦时 | {metrics['power_usage_growth_pct']:.1f}% |",
            "",
            "### 财务与合规影响评估",
            (
                f"2025 年新增收入规模为 {metrics['revenue_2025_rmb'] - metrics['revenue_2024_rmb']} 元，但同期用电量仅增加 "
                f"{metrics['power_consumption_2025_kwh'] - metrics['power_consumption_2024_kwh']} 千瓦时。"
                "若无充分的外协、产能外移或跨期确认说明，则收入增长的真实性、可持续性及第一还款来源稳定性将受到影响。"
            ),
            "",
            "### 补充尽调建议",
            "建议补充外协加工合同、结算单、仓储入库单、验收单及核心客户回款凭证。",
            "",
            "## 第三章 财务与会计主要数据指标评析",
            "### 风险定性",
            (
                f"第一大客户 {metrics['largest_customer_name']} 贡献收入 {metrics['largest_customer_revenue_rmb']} 元，占比 "
                f"{metrics['largest_customer_revenue_share_pct']:.1f}%。同时，2025 年报表收入高于可识别回款及税务申报收入，当前财务质量判断仍需以勾稽核验结果为基础"
                f" [{source_prefix}/financial_statements.md] [{source_prefix}/bank_statement_summary.md] [{source_prefix}/tax_filing_summary.md]。"
            ),
            "",
            "### 核查证据",
            "| 指标 | 数值 | 关注点 |",
            "| --- | --- | --- |",
            f"| 2025 年营业收入 | {metrics['revenue_2025_rmb']} 元 | 需核验真实性与持续性 |",
            f"| 2025 年可识别回款 | {metrics['bank_receipts_2025_rmb']} 元 | 低于报表收入 |",
            f"| 2025 年税务申报收入 | {metrics['tax_filed_revenue_2025_rmb']} 元 | 低于报表收入 |",
            f"| 第四季度财务确认收入 | {pack['q4_booked_revenue_rmb']} 元 | 需与跨期回款对照 |",
            f"| 第四季度银行实际回款 | {pack['q4_receipts_rmb']} 元 | 与报表收入存在差异 |",
            "",
            "| 勾稽维度 | 财务口径 | 外部/交叉口径 | 差异说明 |",
            "| --- | --- | --- | --- |",
            f"| 收入与回款 | {metrics['revenue_2025_rmb']} 元 | {metrics['bank_receipts_2025_rmb']} 元 | 差额 {metrics['revenue_2025_rmb'] - metrics['bank_receipts_2025_rmb']} 元 |",
            f"| 收入与税票 | {metrics['revenue_2025_rmb']} 元 | {metrics['tax_filed_revenue_2025_rmb']} 元 | 差额 {metrics['revenue_2025_rmb'] - metrics['tax_filed_revenue_2025_rmb']} 元 |",
            "",
            "### 财务与合规影响评估",
            (
                f"{cash_tax_issue.get('problem_definition', '当前收入、回款及税票之间存在勾稽差异')}。"
                "如差异来源于跨期确认、未回款应收集中或申报时点差异，应提供完整桥接解释；否则将影响对收入质量、现金转化能力和还款来源真实性的判断。"
            ),
            "",
            "### 补充尽调建议",
            "建议补充税会差异桥接表、2026 年 1 月回款对账单、主要客户回款台账及收入确认政策说明。",
            "",
            "## 第四章 资产结构穿透与资金独立性分析",
            "### 风险定性",
            (
                f"除收入真实性外，核心客户 {metrics['largest_customer_name']} 与实控人家庭控制关系重合，说明交易独立性与主体独立性存在交叉风险。"
                "若关联客户交易占比较高但未单列披露，将放大资金回收、定价公允性和关联占用识别难度"
                f" [{source_prefix}/financial_statements.md] [{source_prefix}/shareholding_register.md] [{source_prefix}/board_meeting_minutes_secret.md]。"
            ),
            "",
            "### 核查证据",
            "| 资产/资金事项 | 当前情况 | 关注点 | 风险结论 |",
            "| --- | --- | --- | --- |",
            f"| 第一大客户收入占比 | {metrics['largest_customer_revenue_share_pct']:.1f}% | 集中度偏高且存在关联控制重合 | 持续关注 |",
            f"| 第一大客户收入金额 | {metrics['largest_customer_revenue_rmb']} 元 | 需复核交易公允性与独立性 | 持续关注 |",
            "",
            "### 财务与合规影响评估",
            (
                f"{related_issue.get('problem_definition', '关联交易和控制链条存在交叉风险')}。"
                "若核心客户兼具关联属性，相关收入质量、回款独立性及未来持续合作稳定性均需在授信判断中从严审视。"
            ),
            "",
            "### 补充尽调建议",
            "建议补充关联交易披露说明、定价依据、独立性承诺及关联资金往来说明。",
            "",
            "## 第五章 有息债务期限结构与表外负债排查",
            "### 风险定性",
            (
                f"董事会纪要显示，公司已批准为兄弟企业 {pack['guaranteed_company_name']} 提供 {metrics['undisclosed_guarantee_amount_rmb']} 元全额连带责任担保，"
                "且要求财务不对外披露。该事项已超出一般日常经营事项范围，属于需要单独识别的重大或有负债风险"
                f" [{source_prefix}/board_meeting_minutes_secret.md]。"
            ),
            "",
            "### 核查证据",
            "| 债务/担保事项 | 涉及对象 | 金额 | 是否披露 | 核查结论 |",
            "| --- | --- | --- | --- | --- |",
            f"| 连带责任担保 | {pack['guaranteed_company_name']} | {metrics['undisclosed_guarantee_amount_rmb']} 元 | 否 | 重大关注 |",
            "",
            "### 财务与合规影响评估",
            (
                f"{liability_issue.get('problem_definition', '存在未披露表外连带责任风险')}。"
                "被担保主体已出现逾期迹象，意味着该担保可能从潜在责任向实际代偿压力转化，并对短期偿债能力、资产保全及诉讼风险形成冲击"
                f" [{source_prefix}/litigation_and_penalties.md]。"
            ),
            "",
            "### 补充尽调建议",
            "建议立即补充担保合同原件、董事会决议原件、被担保企业债务情况及后续整改计划。",
            "",
            "## 第六章 初步尽调结论与补件建议",
            "### 风险定性",
            (
                f"本次审阅识别出 {high_count} 项高风险事项"
                f"{f'及 {medium_count} 项中风险事项' if medium_count else ''}。"
                f"{revenue_issue.get('finding_title', '收入真实性风险')}、{related_issue.get('finding_title', '关联交易风险')}、"
                f"{liability_issue.get('finding_title', '表外负债风险')} 同时存在，说明经营真实性、治理独立性和债务完整性三个维度出现联动问题。"
            ),
            "",
            "### 核查证据",
            "| 待补材料 | 影响章节 | 影响判断 | 紧急程度 |",
            "| --- | --- | --- | --- |",
            "| 外协合同与结算单 | 第二章 | 收入真实性与业务闭环 | 高 |",
            "| 回款台账与税会桥接表 | 第三章 | 现金回收质量与财务勾稽 | 高 |",
            "| 关联交易披露与定价说明 | 第四章 | 主体独立性与交易公允性 | 高 |",
            "| 担保合同及整改材料 | 第五章 | 或有负债与偿债能力 | 高 |",
            "",
            "### 财务与合规影响评估",
            "在关键资料尚未补齐前，现阶段难以形成无保留正面结论。若上述风险未能被充分解释或缓释，可能直接影响授信准入、额度安排与缓释条件设计。",
            "",
            "### 补充尽调建议",
            (
                "建议补件后再审。建议先完成外协与回款补证、关联交易专项说明、担保敞口核验及管理层补充说明，再进入下一轮人工复核"
                f" [{source_prefix}/management_interview_notes.md] [{source_prefix}/bank_statement_summary.md] [{source_prefix}/board_meeting_minutes_secret.md]。"
            ),
            "",
            f"版本标识：V{version_number}",
        ]
    )


def extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n".join(chunk.strip() for chunk in chunks if chunk.strip())


def infer_file_text(path: Path, fallback_text: str = "") -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return extract_pdf_text(path) or fallback_text
        if suffix == ".docx":
            return extract_docx_text(path) or fallback_text
        return path.read_text(encoding="utf-8", errors="ignore") or fallback_text
    except Exception:
        return fallback_text


def content_type_for_suffix(suffix: str) -> str:
    return (
        mimetypes.types_map.get(f".{suffix.lower()}")
        or "application/octet-stream"
    )


def knowledge_file_record(file_item: dict) -> dict:
    return {
        **file_item,
        "download_url": f"/api/knowledge-files/{quote(file_item['id'])}/download",
        "view_url": f"/api/knowledge-files/{quote(file_item['id'])}/view",
    }


def render_preview_html(title: str, description: str, content_text: str) -> bytes:
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
  <head>
    <meta charset="UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0" />
    <title>{html_escape(title)}</title>
    <style>
      body {{
        margin: 0;
        font-family: "PingFang SC", "Microsoft YaHei", sans-serif;
        background: #f4f7ef;
        color: #123020;
      }}
      .page {{
        max-width: 920px;
        margin: 32px auto;
        padding: 28px 32px;
        background: #fff;
        border-radius: 16px;
        box-shadow: 0 18px 36px rgba(10,17,12,0.08);
      }}
      h1 {{
        margin: 0 0 12px;
        font-size: 32px;
      }}
      p {{
        line-height: 1.8;
        color: #5f6f65;
      }}
      pre {{
        white-space: pre-wrap;
        line-height: 1.85;
        font-family: "PingFang SC", "Microsoft YaHei", sans-serif;
        color: #1b3f2b;
        background: #f6f9f0;
        border: 1px solid #d7e5ca;
        border-radius: 12px;
        padding: 18px;
      }}
    </style>
  </head>
  <body>
    <div class="page">
      <h1>{html_escape(title)}</h1>
      <p>{html_escape(description)}</p>
      <pre>{html_escape(content_text or "当前文件暂无可预览文本，请下载后查看原文件。")}</pre>
    </div>
  </body>
</html>"""
    return html.encode("utf-8")


def call_deepseek_chat(
    messages: list[dict],
    *,
    system_prompt: str | None = None,
    model: str | None = None,
) -> dict:
    api_key = get_deepseek_api_key()
    if not api_key:
        raise RuntimeError("DeepSeek API Key 未配置")

    model_name = (model or DEEPSEEK_MODEL).strip() or "deepseek-v4-flash"
    payload_messages = []
    if system_prompt:
        payload_messages.append({"role": "system", "content": system_prompt})
    payload_messages.extend(messages)
    payload = json.dumps(
        {
            "model": model_name,
            "temperature": 0.2,
            "max_tokens": 12000,
            "thinking": {"type": "disabled"},
            "messages": payload_messages,
        }
    ).encode("utf-8")
    request = Request(
        "https://api.deepseek.com/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    ssl_context = ssl.create_default_context(cafile=certifi.where())
    with urlopen(request, timeout=420, context=ssl_context) as response:
        return json.loads(response.read().decode("utf-8"))


def build_system_admin_payload() -> dict:
    return {
        "users": [
            {"name": "王晓明", "department": "公司金融部", "job_title": "客户经理", "status": "启用", "role": "客户经理"},
            {"name": "李敏", "department": "风险管理部", "job_title": "风险经理", "status": "启用", "role": "风险审查员"},
            {"name": "陈卓", "department": "运营管理部", "job_title": "系统管理员", "status": "启用", "role": "系统管理员"},
        ],
        "roles": [
            {"name": "客户经理", "description": "负责企业资料整理、尽调材料上传、报告发起", "menus": "风险视图、全流程引擎、知识库"},
            {"name": "风险审查员", "description": "负责风险视图判断、报告预审和审批建议", "menus": "风险视图、全流程引擎、知识库"},
            {"name": "系统管理员", "description": "负责用户权限、知识权限和菜单配置", "menus": "系统管理、知识库"},
        ],
        "permissions": [
            {"name": "页面权限", "description": "控制顶部主栏与左侧菜单访问", "scope": "按角色"},
            {"name": "按钮权限", "description": "控制生成报告、预审、上传、下载等动作", "scope": "按角色"},
            {"name": "数据范围", "description": "控制可查看的企业、案件和知识文件范围", "scope": "按部门/角色"},
        ],
        "menus": [
            {"name": "顶部主栏", "description": "风险视图 / 全流程引擎 / 知识库 / 系统管理", "visibility": "全部启用"},
            {"name": "风险视图菜单", "description": "8 个客户风险专题页", "visibility": "启用"},
            {"name": "全流程引擎菜单", "description": "6 个授信流程节点", "visibility": "启用"},
        ],
        "knowledge_permissions": [
            {"role": "客户经理", "laws": "查看", "policies": "查看/引用", "experience": "查看/引用"},
            {"role": "风险审查员", "laws": "查看/引用", "policies": "查看/引用", "experience": "查看/引用"},
            {"role": "系统管理员", "laws": "上传/查看/引用", "policies": "上传/查看/引用", "experience": "上传/查看/引用"},
        ],
        "logs": [
            {"time": "2026-05-09 09:20", "user": "王晓明", "action": "上传尽调资料", "target": "比亚迪股份有限公司 / 财务与经营材料"},
            {"time": "2026-05-09 10:05", "user": "李敏", "action": "执行报告预审", "target": "尽调报告 V2"},
            {"time": "2026-05-09 10:40", "user": "陈卓", "action": "调整知识库权限", "target": "专家经验分类"},
        ],
    }


def ensure_knowledge_runtime() -> dict:
    if KNOWLEDGE_BASE_RUNTIME["initialized"]:
        return KNOWLEDGE_BASE_RUNTIME

    ensure_knowledge_library_dirs()
    if KNOWLEDGE_META_PATH.exists():
        payload = json.loads(KNOWLEDGE_META_PATH.read_text(encoding="utf-8"))
        files = payload.get("files", [])
        next_id = payload.get("next_id", len(files) + 1)
    else:
        files, next_id = create_seeded_knowledge_files()

    KNOWLEDGE_BASE_RUNTIME["initialized"] = True
    KNOWLEDGE_BASE_RUNTIME["next_id"] = next_id
    KNOWLEDGE_BASE_RUNTIME["files"] = files
    return KNOWLEDGE_BASE_RUNTIME


def knowledge_file_map() -> dict[str, dict]:
    runtime = ensure_knowledge_runtime()
    return {item["id"]: item for item in runtime["files"]}


def serialize_knowledge_files() -> dict:
    runtime = ensure_knowledge_runtime()
    return {
        "categories": [
            {"id": key, "title": value}
            for key, value in KNOWLEDGE_CATEGORY_TITLES.items()
        ],
        "files": [knowledge_file_record(item) for item in runtime["files"]],
    }


def ensure_material_runtime(company_code: str) -> dict:
    runtime = DUE_DILIGENCE_MATERIALS_RUNTIME.get(company_code)
    if runtime:
        return runtime

    runtime = {
        "next_id": 1,
        "files": [],
    }
    DUE_DILIGENCE_MATERIALS_RUNTIME[company_code] = runtime
    return runtime


def ensure_pdf_font() -> str:
    if not REPORTLAB_AVAILABLE:
        return "Helvetica"
    if PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return PDF_FONT_NAME
    for path in PDF_FONT_PATHS:
        if path.exists():
            pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(path)))
            return PDF_FONT_NAME
    return "Helvetica"


def split_text_lines(text: str, chunk_size: int = 26) -> list[str]:
    normalized = str(text or "").replace("\r", "")
    lines: list[str] = []
    for raw_line in normalized.split("\n"):
        raw_line = raw_line.strip()
        if not raw_line:
            lines.append("")
            continue
        start = 0
        while start < len(raw_line):
            lines.append(raw_line[start : start + chunk_size])
            start += chunk_size
    return lines


def build_pdf_bytes(title: str, version_label: str, full_text: str) -> bytes:
    if not REPORTLAB_AVAILABLE:
        return full_text.encode("utf-8")

    font_name = ensure_pdf_font()
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin_x = 54
    margin_right = width - 54
    cursor_y = height - 72

    def draw_page_header(page_title: str) -> float:
        pdf.setStrokeColorRGB(0.84, 0.9, 0.8)
        pdf.setLineWidth(1)
        pdf.line(margin_x, height - 52, margin_right, height - 52)
        pdf.setFont(font_name, 10)
        pdf.setFillColorRGB(0.37, 0.46, 0.4)
        pdf.drawString(margin_x, height - 44, page_title)
        pdf.drawRightString(margin_right, height - 44, f"第 {pdf.getPageNumber()} 页")
        pdf.setFillColorRGB(0.07, 0.19, 0.13)
        return height - 84

    def new_page(page_title: str = "尽职调查报告") -> float:
        pdf.showPage()
        return draw_page_header(page_title)

    def wrap_lines(text: str, max_chars: int = 36) -> list[str]:
        output: list[str] = []
        for raw_line in text.replace("\r", "").split("\n"):
            raw_line = raw_line.strip()
            if not raw_line:
                output.append("")
                continue
            start = 0
            while start < len(raw_line):
                output.append(raw_line[start : start + max_chars])
                start += max_chars
        return output

    def draw_paragraph(text: str, *, font_size: int = 11, leading: int = 20) -> float:
        nonlocal cursor_y
        pdf.setFont(font_name, font_size)
        for line in wrap_lines(text, 38):
            if cursor_y < 72:
                cursor_y = new_page()
                pdf.setFont(font_name, font_size)
            if line == "":
                cursor_y -= leading // 2
                continue
            pdf.drawString(margin_x, cursor_y, line)
            cursor_y -= leading
        return cursor_y

    def draw_table(table_lines: list[str]) -> float:
        nonlocal cursor_y
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            rows.append(cells)
        if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in rows[1]):
            rows.pop(1)
        if not rows:
            return cursor_y

        col_count = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (col_count - len(row)) for row in rows]
        table_width = margin_right - margin_x
        col_width = table_width / max(col_count, 1)
        font_size = 9
        line_height = 13

        for row_index, row in enumerate(normalized_rows):
            wrapped_cells = [
                wrap_lines(cell, max(6, int((col_width - 18) / max(font_size * 0.9, 1)))) or [""]
                for cell in row
            ]
            row_height = max(len(cell_lines) for cell_lines in wrapped_cells) * line_height + 12
            if cursor_y - row_height < 72:
                cursor_y = new_page("尽职调查报告")
            for col_index, cell_lines in enumerate(wrapped_cells):
                x = margin_x + col_index * col_width
                y = cursor_y
                if row_index == 0:
                    pdf.setFillColorRGB(0.95, 0.97, 0.92)
                    pdf.rect(x, y - row_height, col_width, row_height, fill=1, stroke=0)
                pdf.setStrokeColorRGB(0.82, 0.88, 0.78)
                pdf.rect(x, y - row_height, col_width, row_height, fill=0, stroke=1)
                pdf.setFillColorRGB(0.07, 0.19, 0.13 if row_index == 0 else 0.22)
                pdf.setFont(font_name, font_size)
                text_y = y - 10
                for cell_line in cell_lines:
                    pdf.drawString(x + 8, text_y, cell_line)
                    text_y -= line_height
            cursor_y -= row_height
        cursor_y -= 10
        pdf.setFillColorRGB(0.22, 0.29, 0.24)
        return cursor_y

    def draw_markdown_content(text: str) -> float:
        nonlocal cursor_y
        lines = text.replace("\r", "").split("\n")
        paragraph: list[str] = []
        table_lines: list[str] = []
        list_items: list[str] = []
        ordered_list = False

        def flush_paragraph() -> None:
            nonlocal cursor_y, paragraph
            if not paragraph:
                return
            draw_paragraph(" ".join(paragraph), font_size=11, leading=20)
            cursor_y -= 4
            paragraph = []

        def flush_table() -> None:
            nonlocal table_lines
            if not table_lines:
                return
            draw_table(table_lines)
            table_lines = []

        def flush_list() -> None:
            nonlocal cursor_y, list_items
            if not list_items:
                return
            for index, item in enumerate(list_items, start=1):
                marker = f"{index}. " if ordered_list else "• "
                draw_paragraph(f"{marker}{item}", font_size=11, leading=20)
            cursor_y -= 4
            list_items = []

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                flush_paragraph()
                flush_list()
                flush_table()
                continue
            if line.startswith("|"):
                flush_paragraph()
                flush_list()
                table_lines.append(line)
                continue
            flush_table()

            if line.startswith("### "):
                flush_paragraph()
                flush_list()
                if cursor_y < 100:
                    cursor_y = new_page("尽职调查报告")
                pdf.setFont(font_name, 13)
                pdf.setFillColorRGB(0.07, 0.19, 0.13)
                pdf.drawString(margin_x, cursor_y, line.replace("### ", "", 1))
                cursor_y -= 24
                pdf.setFillColorRGB(0.22, 0.29, 0.24)
                continue
            if line.startswith("#### "):
                flush_paragraph()
                flush_list()
                if cursor_y < 92:
                    cursor_y = new_page("尽职调查报告")
                pdf.setFont(font_name, 12)
                pdf.setFillColorRGB(0.16, 0.25, 0.19)
                pdf.drawString(margin_x, cursor_y, line.replace("#### ", "", 1))
                cursor_y -= 22
                pdf.setFillColorRGB(0.22, 0.29, 0.24)
                continue

            ordered_match = re.match(r"^(\d+)\.\s+(.+)$", line)
            unordered_match = re.match(r"^[-*]\s+(.+)$", line)
            if ordered_match:
                flush_paragraph()
                if list_items and not ordered_list:
                    flush_list()
                ordered_list = True
                list_items.append(ordered_match.group(2))
                continue
            if unordered_match:
                flush_paragraph()
                if list_items and ordered_list:
                    flush_list()
                ordered_list = False
                list_items.append(unordered_match.group(1))
                continue

            flush_list()
            paragraph.append(line)

        flush_paragraph()
        flush_list()
        flush_table()
        return cursor_y

    pdf.setTitle(f"{title}-{version_label}")
    pdf.setFillColorRGB(0.07, 0.19, 0.13)
    pdf.setFont(font_name, 24)
    pdf.drawString(margin_x, cursor_y, title)
    cursor_y -= 42
    pdf.setFont(font_name, 15)
    pdf.drawString(margin_x, cursor_y, "尽职调查报告")
    cursor_y -= 28
    pdf.setFont(font_name, 12)
    pdf.setFillColorRGB(0.37, 0.46, 0.4)
    pdf.drawString(margin_x, cursor_y, f"版本：{version_label}")
    cursor_y -= 22
    pdf.drawString(margin_x, cursor_y, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    cursor_y -= 36
    pdf.setFillColorRGB(0.07, 0.19, 0.13)
    draw_paragraph("本报告基于企业结构化数据、补充尽调材料、知识库引用文件及公开信息补全结果生成，用于银行内部授信尽职调查参考。", font_size=12, leading=22)
    cursor_y -= 12

    section_heading_pattern = re.compile(r"^(?:第[一二三四五六七八九十百0-9]+章|[一二三四五六七八九十]+、)")
    sections = []
    current_title = None
    current_lines: list[str] = []
    normalized_lines = full_text.replace("\r", "").split("\n")
    for raw_line in normalized_lines:
        line = raw_line.strip()
        if not line:
            current_lines.append("")
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            line = line.replace("## ", "", 1).strip()
        if section_heading_pattern.match(line):
            if current_title:
                sections.append((current_title, "\n".join(current_lines).strip()))
            current_title = line
            current_lines = []
        else:
            current_lines.append(line)
    if current_title:
        sections.append((current_title, "\n".join(current_lines).strip()))
    if not sections and full_text.strip():
        sections = [("尽职调查报告正文", full_text.strip())]

    if sections:
        cursor_y = new_page("目录")
        pdf.setFont(font_name, 18)
        pdf.setFillColorRGB(0.07, 0.19, 0.13)
        pdf.drawString(margin_x, cursor_y, "目录")
        cursor_y -= 30
        pdf.setFont(font_name, 11)
        for index, (section_title, _) in enumerate(sections, start=1):
            if cursor_y < 72:
                cursor_y = new_page("目录")
                pdf.setFont(font_name, 11)
            pdf.drawString(margin_x, cursor_y, f"{index}. {section_title}")
            cursor_y -= 20

    if sections:
        cursor_y = new_page("尽职调查报告")
        for section_title, section_body in sections:
            if cursor_y < 96:
                cursor_y = new_page("尽职调查报告")
            pdf.setFont(font_name, 16)
            pdf.setFillColorRGB(0.07, 0.19, 0.13)
            pdf.drawString(margin_x, cursor_y, section_title)
            cursor_y -= 26
            pdf.setFillColorRGB(0.22, 0.29, 0.24)
            draw_markdown_content(section_body)
            cursor_y -= 10

    pdf.save()
    return buffer.getvalue()


def extract_upload_preview(field: FieldStorage) -> str:
    if not getattr(field, "file", None):
        return ""
    content = field.file.read()
    if isinstance(content, str):
        text = content
    else:
        text = content.decode("utf-8", errors="ignore")
    field.file.seek(0)
    return compact_text(text, 240)


def find_version(company_code: str, version_id: str) -> dict | None:
    runtime = REPORT_RUNTIME.get(company_code, {})
    for version in runtime.get("versions", []):
        if version["id"] == version_id:
            return version
    return None


def serialize_review_result(review_result: dict | None) -> dict | None:
    if not review_result:
        return None
    return {
        "review_status": review_result["review_status"],
        "overall_result": review_result["overall_result"],
        "reviewed_at": review_result["reviewed_at"],
        "counts": review_result["counts"],
        "review_summary": review_result.get("review_summary", ""),
        "findings": review_result["findings"],
    }


def serialize_report_version(company_code: str, version: dict) -> dict:
    knowledge_map = knowledge_file_map()
    return {
        "id": version["id"],
        "version_label": version["version_label"],
        "created_at": version["created_at"],
        "status": version["status"],
        "based_on": version["based_on"],
        "section_list": version["section_list"],
        "full_text": version["full_text"],
        "review_files": [
            {
                "id": item["id"],
                "name": item["name"],
                "category": item["category"],
                "uploaded_at": item["uploaded_at"],
                "size": item["size"],
            }
            for item in version["review_files"]
        ],
        "knowledge_files": [
            knowledge_file_record(knowledge_map[file_id])
            for file_id in version.get("knowledge_file_ids", [])
            if file_id in knowledge_map
        ],
        "selected_data_source_ids": version.get("selected_data_source_ids", get_all_report_data_source_ids()),
        "references": version.get("references") or [],
        "review_result": serialize_review_result(version.get("review_result")),
        "generation_mode": version.get("generation_mode", "template"),
        "review_edit_text": version.get("review_edit_text", version.get("full_text", "")),
        "review_saved_at": version.get("review_saved_at"),
        "pdf_url": f"/api/company/{quote(company_code)}/report-versions/{quote(version['id'])}/pdf",
    }


def build_fallback_report_sections(detail: dict, version_number: int) -> tuple[list[dict], str, list[str]]:
    company = detail["company"]
    case_info = detail.get("case_info") or {}
    recommendation = detail.get("recommendation") or {}
    metrics = metric_map(detail.get("latest_metrics") or [])
    top_findings = detail.get("validation_findings") or []
    guarantee = (detail.get("guarantees") or [None])[0]

    sections = [
        {
            "id": "basic-history",
            "title": "第一章 基本情况与历史沿革核查",
            "content": (
                "### 风险定性\n"
                f"{company['name']}位于 {company.get('region_province') or '—'}{company.get('region_city') or ''}，所属行业为 "
                f"{company.get('subindustry') or company.get('industry_category') or '—'}。根据当前已获取资料，主体基础信息已同步，但历史沿革、治理结构和实控链条仍需进一步核实。\n\n"
                "### 核查证据\n"
                f"| 项目 | 当前信息 | 核查说明 |\n| --- | --- | --- |\n| 企业名称 | {company['name']} | 已同步 |\n"
                f"| 所在地区 | {(company.get('region_province') or '—')}{company.get('region_city') or ''} | 可继续补充工商底稿 |\n"
                f"| 所属行业 | {company.get('subindustry') or company.get('industry_category') or '—'} | 可继续结合公开信息补全 |\n\n"
                "### 财务与合规影响评估\n主体沿革和治理信息若存在缺口，可能影响授信准入与独立性判断。\n\n"
                "### 补充尽调建议\n建议补充工商档案、章程、股权变更与实际控制人说明。"
            ),
        },
        {
            "id": "operation-analysis",
            "title": "第二章 主营业务与经营模式深挖",
            "content": (
                "### 风险定性\n"
                f"企业当前经营状态为 {company.get('operating_status') or '—'}，授信申请产品为 {case_info.get('product_type') or '—'}，"
                f"第一还款来源为 {case_info.get('primary_repayment_source') or '—'}。但业务闭环仍需结合合同、订单、发票与回款链路核验。\n\n"
                "### 核查证据\n"
                f"| 维度 | 当前情况 | 风险提示 |\n| --- | --- | --- |\n| 合同数据 | {len(detail.get('contracts') or [])} 份 | 需验证真实性 |\n"
                f"| 订单数据 | {len(detail.get('orders') or [])} 条 | 需验证履约情况 |\n"
                f"| 发票数据 | {len(detail.get('invoices') or [])} 张 | 需验证与回款匹配度 |\n\n"
                "### 财务与合规影响评估\n业务闭环支撑不足时，将削弱主营业务真实性和第一还款来源稳定性的判断。\n\n"
                "### 补充尽调建议\n建议补充核心客户合同、验收单和回款凭证。"
            ),
        },
        {
            "id": "financial-analysis",
            "title": "第三章 财务与会计主要数据指标评析",
            "content": (
                "### 风险定性\n"
                f"最新口径下营业收入 {format_currency(metrics.get('revenue'))}，净利润 {format_currency(metrics.get('net_profit'))}，货币资金 {format_currency(metrics.get('cash'))}。"
                f"应收账款 {format_currency(metrics.get('accounts_receivable'))}，存货 {format_currency(metrics.get('inventory'))}，需关注利润与现金回收是否匹配。\n\n"
                "### 核查证据\n"
                f"| 指标 | 当前值 | 关注点 |\n| --- | --- | --- |\n| 营业收入 | {format_currency(metrics.get('revenue'))} | 需核验真实性 |\n"
                f"| 净利润 | {format_currency(metrics.get('net_profit'))} | 需结合现金流判断 |\n"
                f"| 货币资金 | {format_currency(metrics.get('cash'))} | 需关注可支配性 |\n"
                f"| 应收账款 | {format_currency(metrics.get('accounts_receivable'))} | 需关注账龄 |\n"
                f"| 存货 | {format_currency(metrics.get('inventory'))} | 需关注周转 |\n\n"
                "### 财务与合规影响评估\n若收入、利润与现金流偏离，或应收、存货上升明显，将影响对偿债能力和经营质量的判断。\n\n"
                "### 补充尽调建议\n建议补充近三年报表明细、流水、税票及往来账龄分析。"
            ),
        },
        {
            "id": "asset-independence",
            "title": "第四章 资产结构穿透与资金独立性分析",
            "content": (
                "### 风险定性\n"
                f"当前货币资金为 {format_currency(metrics.get('cash'))}，应收账款为 {format_currency(metrics.get('accounts_receivable'))}，存货为 {format_currency(metrics.get('inventory'))}。"
                "若存在资金受限、关联占用或资产质量偏弱情形，将直接影响流动性判断。\n\n"
                "### 核查证据\n"
                f"| 资产/资金事项 | 当前情况 | 关注点 | 风险结论 |\n| --- | --- | --- | --- |\n| 货币资金 | {format_currency(metrics.get('cash'))} | 可支配性待核实 | 持续关注 |\n"
                f"| 应收账款 | {format_currency(metrics.get('accounts_receivable'))} | 账龄待补充 | 持续关注 |\n"
                f"| 存货 | {format_currency(metrics.get('inventory'))} | 周转与跌价待补充 | 持续关注 |\n\n"
                "### 财务与合规影响评估\n资产质量和资金独立性若存在问题，将削弱短期偿债保障和主体独立性判断。\n\n"
                "### 补充尽调建议\n建议补充资产权属、盘点记录和银行受限说明。"
            ),
        },
        {
            "id": "debt-guarantee",
            "title": "第五章 有息债务期限结构与表外负债排查",
            "content": (
                "### 风险定性\n"
                f"当前授信缓释要求为 {recommendation.get('guarantee_requirement') or '—'}，已登记担保措施 {len(detail.get('guarantees') or [])} 项。"
                f"重点抵质押物暂识别为 {compact_text((guarantee or {}).get('asset_name') or (guarantee or {}).get('guarantee_type') or '待补充')}，仍需核验可执行性。\n\n"
                "### 核查证据\n"
                f"| 担保/债务事项 | 当前情况 | 风险提示 |\n| --- | --- | --- |\n| 担保措施数量 | {len(detail.get('guarantees') or [])} 项 | 需核验真实性与足值性 |\n"
                f"| 重点押品 | {compact_text((guarantee or {}).get('asset_name') or (guarantee or {}).get('guarantee_type') or '待补充')} | 需核验权属、评估值与处置性 |\n\n"
                "### 财务与合规影响评估\n担保、押品或潜在或有负债披露不充分时，将影响第二还款来源判断和缓释有效性。\n\n"
                "### 补充尽调建议\n建议补充借款合同、担保合同、押品权属及评估材料。"
            ),
        },
        {
            "id": "risk-conclusion",
            "title": "第六章 初步尽调结论与补件建议",
            "content": (
                "### 风险定性\n"
                f"当前风险分层为 {company.get('risk_tier') or '—'}，授信建议状态为 {recommendation.get('recommendation_status') or '—'}。"
                f"核验发现共 {len(top_findings)} 条，需重点关注 {compact_text(top_findings[0].get('finding_title') if top_findings else '资料完整性与关键风险点仍待补充')}。\n\n"
                "### 核查证据\n"
                "| 待补材料 | 影响章节 | 影响判断 | 紧急程度 |\n| --- | --- | --- | --- |\n"
                "| 主体与治理补充材料 | 第一章 | 主体独立性与治理稳定性 | 高 |\n"
                "| 核心业务闭环与回款材料 | 第二、三章 | 经营真实性与还款来源 | 高 |\n"
                "| 担保与押品完整材料 | 第五章 | 风险缓释有效性 | 中 |\n\n"
                "### 财务与合规影响评估\n关键资料未补齐前，当前尽调结论仍应保持审慎。\n\n"
                "### 补充尽调建议\n建议补件后再审，并结合后续内外规预审结果形成最终审批意见。"
            ),
        },
    ]

    based_on = [
        "基本信息、经营情况、财务情况已同步",
        "行业与外部环境信息已纳入章节",
        "抵质押物与缓释措施已形成摘要",
        "可继续通过内外规预审形成修改闭环",
    ]

    full_text = "\n\n".join(
        [f"尽职调查报告 {version_number} 版"]
        + [f"{index + 1}. {section['title']}\n{section['content']}" for index, section in enumerate(sections)]
    )
    return sections, full_text, based_on


REPORT_SYSTEM_PROMPT = (
    "你是银行公司授信尽职调查报告撰写智能体。你必须以正式、审慎、客观、完整的中文输出报告。"
    "不得虚构事实，缺失信息要明确写待补充核实。报告必须章节完整、前后判断一致、风险与结论呼应。"
)
SKILL_PIPELINE_REPORT_SYSTEM_PROMPT = (
    "你是一名银行授信尽职调查报告主笔。"
    "你必须严格遵循用户提供的 auditing-data 技能、writing-report 技能和 Markdown 模板要求输出正式尽调报告。"
    "你输出的必须是纯中文 Markdown 长文本，结构完整，表格表头为中文，判断审慎，结论与证据一致。"
    "不得输出解释性前言，不得输出 JSON，不得忽略高风险 findings。"
)
REPORT_MARKDOWN_TEMPLATE = """# 尽职调查报告

## 第一章 基本情况与历史沿革核查
### 风险定性
### 核查证据
### 财务与合规影响评估
### 补充尽调建议

## 第二章 主营业务与经营模式深挖
### 风险定性
### 核查证据
### 财务与合规影响评估
### 补充尽调建议

## 第三章 财务与会计主要数据指标评析
### 风险定性
### 核查证据
### 财务与合规影响评估
### 补充尽调建议

## 第四章 资产结构穿透与资金独立性分析
### 风险定性
### 核查证据
### 财务与合规影响评估
### 补充尽调建议

## 第五章 有息债务期限结构与表外负债排查
### 风险定性
### 核查证据
### 财务与合规影响评估
### 补充尽调建议

## 第六章 初步尽调结论与补件建议
### 风险定性
### 核查证据
### 财务与合规影响评估
### 补充尽调建议
"""
REVIEW_SYSTEM_PROMPT = (
    "你是银行公司授信尽职调查预审智能体。你需要基于尽调报告正文、企业结构化数据、外部法规、内部制度和专家经验，"
    "输出严谨、可执行、面向客户经理的风险提示清单。不得虚构未提供的依据，不得忽略已经提供的文件。"
)


def json_block(name: str, payload: object) -> str:
    return f"{name}\n```json\n{json.dumps(payload, ensure_ascii=False, indent=2)}\n```"


def get_all_report_data_source_ids() -> list[str]:
    return [item["id"] for group in REPORT_DATA_SOURCE_GROUPS for item in group["items"]]


def resolve_selected_data_source_ids(selected_ids: list[str] | None) -> list[str]:
    valid_ids = set(get_all_report_data_source_ids())
    filtered = [item for item in (selected_ids or []) if item in valid_ids]
    return filtered or get_all_report_data_source_ids()


def summarize_selected_data_sources(selected_ids: list[str]) -> list[dict]:
    selected = set(resolve_selected_data_source_ids(selected_ids))
    groups: list[dict] = []
    for group in REPORT_DATA_SOURCE_GROUPS:
        picked = [item for item in group["items"] if item["id"] in selected]
        if picked:
            groups.append({"group_title": group["title"], "items": [item["label"] for item in picked]})
    return groups


def build_reference_catalog(
    detail: dict,
    company_code: str,
    knowledge_files: list[dict],
    selected_data_source_ids: list[str] | None = None,
) -> list[dict]:
    refs: list[dict] = []
    seen: set[tuple[str, str, str]] = set()

    def add_ref(
        source_type: str,
        title: str,
        locator: str = "",
        excerpt: str = "",
        source_path: str = "",
        source_id: str = "",
    ) -> None:
        key = (source_type, source_path or source_id or title, locator)
        if key in seen:
            return
        seen.add(key)
        refs.append(
            {
                "id": f"R{len(refs) + 1}",
                "source_type": compact_text(source_type, 60),
                "title": compact_text(title, 120),
                "locator": compact_text(locator, 120),
                "excerpt": compact_text(excerpt, 260),
                "source_path": source_path,
                "source_id": source_id,
            }
        )

    for item in knowledge_files:
        category_title = KNOWLEDGE_CATEGORY_TITLES.get(item.get("category"), "知识库文件")
        locator = "；".join(
            part
            for part in [
                item.get("owner") and f"上传人：{item.get('owner')}",
                item.get("uploaded_at") and f"上传时间：{item.get('uploaded_at')}",
            ]
            if part
        )
        add_ref(
            category_title,
            item.get("name") or "知识库文件",
            locator=locator or "知识库文件",
            excerpt=item.get("description") or item.get("content_text") or "",
            source_path=item.get("stored_name", ""),
            source_id=item.get("id", ""),
        )

    if has_antigravity_mock_pack(company_code):
        audit_summary = build_antigravity_audit_summary(company_code)
        for finding in audit_summary.get("findings", []):
            finding_title = finding.get("finding_title") or "审计发现"
            evidence = "；".join(finding.get("evidence_chain") or [])
            for source in finding.get("source_refs") or []:
                source_path = source.get("file", "")
                source_name = Path(source_path).name
                source_title_map = {
                    "business_license.md": "营业执照资料",
                    "financial_statements.md": "财务报表数据",
                    "electricity_bills_2025.txt": "用电量核验材料",
                    "board_meeting_minutes_secret.md": "董事会会议纪要",
                    "shareholding_register.md": "股东名册与实控关系",
                    "bank_statement_summary.md": "银行流水摘要",
                    "tax_filing_summary.md": "纳税申报摘要",
                    "top_customer_contracts.md": "主要客户合同资料",
                    "management_interview_notes.md": "管理层访谈纪要",
                    "litigation_and_penalties.md": "诉讼处罚材料",
                }
                add_ref(
                    "底稿材料",
                    source_title_map.get(source_name, source_name or source_path or "业务底稿"),
                    locator=source.get("line_hint") or "原始底稿锚点",
                    excerpt=f"{finding_title}：{evidence}",
                    source_path=source_path,
                    source_id=finding_title,
                )

    selected_source_ids = set(resolve_selected_data_source_ids(selected_data_source_ids))
    company = detail.get("company") or {}
    case_info = detail.get("case_info") or {}
    recommendation = detail.get("recommendation") or {}
    metrics = {item.get("metric_code"): item.get("value") for item in detail.get("latest_metrics") or []}
    data_reference_specs = {
        "company": (
            "主体信息表数据",
            "企业主体信息",
            "企业主体基础字段",
            f"企业名称：{company.get('name') or '—'}；成立时间：{company.get('established_on') or '—'}；注册资本：{format_currency(company.get('registered_capital_cny'))}；经营状态：{company.get('operating_status') or '—'}。",
        ),
        "case_info": (
            "授信申请信息",
            "案例与授信申请信息",
            "案件与产品字段",
            f"案例编号：{case_info.get('case_no') or '—'}；申请产品：{case_info.get('product_type') or '—'}；申请金额：{format_currency(case_info.get('requested_amount_cny'))}。",
        ),
        "people": (
            "核心人员表数据",
            "核心人员与角色信息",
            "人员角色明细",
            f"已纳入 {len(detail.get('people') or [])} 条核心人员/担保人/管理层角色记录。",
        ),
        "profile_attributes": (
            "补充画像字段",
            "补充画像属性",
            "客户画像明细",
            f"已纳入 {len(detail.get('profile_attributes') or [])} 条客户画像属性，用于补充主体、经营和管理特征。",
        ),
        "contracts": ("合同数据表", "合同数据", "合同清单", f"已纳入 {len(detail.get('contracts') or [])} 份合同数据，用于核验主营业务和客户关系。"),
        "orders": ("订单数据表", "订单数据", "订单清单", f"已纳入 {len(detail.get('orders') or [])} 条订单数据，用于核验业务闭环。"),
        "invoices": ("发票数据表", "发票数据", "发票清单", f"已纳入 {len(detail.get('invoices') or [])} 张发票数据，用于核验收入与税票一致性。"),
        "bank_summaries": ("回款流水摘要表", "银行摘要与回款", "银行流水摘要", f"已纳入 {len(detail.get('bank_summaries') or [])} 条银行摘要，用于核验回款和现金流。"),
        "related_transactions": ("关联交易明细表", "关联交易摘要", "关联交易记录", f"已纳入 {len(detail.get('related_transactions') or [])} 条关联交易摘要，用于识别关联交易风险。"),
        "latest_metrics": (
            "核心财务指标表数据",
            "最新核心财务指标",
            "财务指标快照",
            f"营业收入：{format_currency(metrics.get('revenue'))}；净利润：{format_currency(metrics.get('net_profit'))}；货币资金：{format_currency(metrics.get('cash'))}；应收账款：{format_currency(metrics.get('accounts_receivable'))}。",
        ),
        "recent_metric_history": ("近年财务趋势表", "近年财务趋势", "年度财务指标", f"已纳入 {len(detail.get('recent_metric_history') or [])} 条近年财务趋势数据。"),
        "reconciliation_checks": ("收入回款税票勾稽表", "勾稽检查结果", "财务勾稽核验", f"已纳入 {len(detail.get('reconciliation_checks') or [])} 条勾稽检查结果，用于核验收入、回款和报表一致性。"),
        "tax_invoice_checks": ("税票一致性核验表", "税票一致性核验", "税票核验记录", f"已纳入 {len(detail.get('tax_invoice_checks') or [])} 条税票一致性核验结果。"),
        "tax_filings": ("纳税申报摘要表", "纳税申报摘要", "税务申报记录", f"已纳入 {len(detail.get('tax_filings') or [])} 条纳税申报摘要。"),
        "credit_history": ("信用历史摘要表", "信用历史摘要", "征信/授信历史记录", f"已纳入 {len(detail.get('credit_history') or [])} 条信用历史记录。"),
        "related_companies": ("关联公司关系表", "关联公司关系", "关联公司清单", f"已纳入 {len(detail.get('related_companies') or [])} 条关联公司关系。"),
        "shareholding_changes": ("股权变更记录表", "股权变更记录", "股权变更清单", f"已纳入 {len(detail.get('shareholding_changes') or [])} 条股权变更记录。"),
        "industry_profile": ("行业画像数据", "行业画像", "行业与政策字段", compact_text((detail.get("industry_profile") or {}).get("policy_direction") or (detail.get("industry_profile") or {}).get("industry_summary"), 220)),
        "peer_comparisons": ("同业比较表", "同业比较", "同业比较记录", f"已纳入 {len(detail.get('peer_comparisons') or [])} 条同业比较数据。"),
        "public_risks": ("公开风险事件表", "公开风险事件", "公开风险记录", f"已纳入 {len(detail.get('public_risks') or [])} 条公开风险事件。"),
        "public_info_enrichment": ("公开信息补全层", "公开信息补全层", "公开资料摘要", "已纳入企业画像、治理摘要、公开风险、行业洞察和年度财务概览等公开信息补全内容。"),
        "guarantees": ("担保与抵质押表", "担保与抵质押信息", "担保/押品记录", f"已纳入 {len(detail.get('guarantees') or [])} 条担保与抵质押记录。"),
        "due_diligence_materials": ("尽调补充材料清单", "尽调补充材料", "客户经理上传材料", f"已纳入 {len(serialize_material_files(company_code))} 份尽调补充材料。"),
        "validation_findings": ("核验发现清单", "核验发现", "系统核验结果", f"已纳入 {len(detail.get('validation_findings') or [])} 条核验发现，用于定位异常和补件要求。"),
        "recommendation": (
            "授信建议与缓释要求",
            "授信建议与缓释要求",
            "授信建议字段",
            f"建议状态：{recommendation.get('recommendation_status') or '—'}；缓释要求：{compact_text(recommendation.get('guarantee_requirement'), 120)}。",
        ),
    }

    for source_id, spec in data_reference_specs.items():
        if source_id not in selected_source_ids:
            continue
        title, source_type, locator, excerpt = spec
        add_ref(
            source_type,
            title,
            locator=locator,
            excerpt=excerpt,
            source_id=source_id,
        )

    if not refs:
        company_name = detail.get("company", {}).get("name") or company_code
        add_ref(
            "报告生成依据",
            f"{company_name}基础尽调数据",
            locator="系统默认输入",
            excerpt="当前版本基于企业结构化数据、补充材料和已选知识文件生成。",
            source_id=company_code,
        )
    return refs


def reference_catalog_prompt_block(references: list[dict]) -> str:
    return json_block("reference_catalog", references)


def reference_marker_aliases(value: object) -> set[str]:
    raw = str(value or "").strip()
    if not raw:
        return set()
    without_dot = raw[2:] if raw.startswith("./") else raw
    aliases = {raw, without_dot}
    if without_dot.startswith(".agent/"):
        aliases.add(f"./{without_dot}")
    if "/" in without_dot:
        aliases.add(Path(without_dot).name)
    return {item for item in aliases if item}


def normalize_reference_markers(text: str, references: list[dict]) -> str:
    normalized = text or ""
    for ref in references:
        marker = f"[{ref['id']}]"
        candidates = [
            ref.get("id", ""),
            ref.get("source_path", ""),
            ref.get("source_id", ""),
            ref.get("title", ""),
        ]
        for candidate in candidates:
            for alias in reference_marker_aliases(candidate):
                normalized = normalized.replace(f"[{alias}]", marker)
    return normalized


def normalize_report_section_references(sections: list[dict], references: list[dict]) -> list[dict]:
    return [
        {
            **section,
            "content": normalize_reference_markers(section.get("content", ""), references),
        }
        for section in sections
    ]


def build_ai_input_payload(detail: dict, company_code: str, knowledge_files: list[dict], selected_data_source_ids: list[str] | None = None) -> dict:
    public_info = load_public_enrichment(company_code)
    selected = set(resolve_selected_data_source_ids(selected_data_source_ids))
    payload = {
        "company": detail.get("company"),
        "case_info": detail.get("case_info"),
        "recommendation": detail.get("recommendation"),
        "people": trim_list(detail.get("people"), 10),
        "related_companies": trim_list(detail.get("related_companies"), 10),
        "shareholding_changes": trim_list(detail.get("shareholding_changes"), 10),
        "profile_attributes": trim_list(detail.get("profile_attributes"), 24),
        "latest_metrics": trim_list(detail.get("latest_metrics"), 12),
        "recent_metric_history": trim_list(detail.get("recent_metric_history"), 12),
        "reconciliation_checks": trim_list(detail.get("reconciliation_checks"), 8),
        "tax_invoice_checks": trim_list(detail.get("tax_invoice_checks"), 6),
        "tax_filings": trim_list(detail.get("tax_filings"), 6),
        "bank_summaries": trim_list(detail.get("bank_summaries"), 6),
        "credit_history": trim_list(detail.get("credit_history"), 8),
        "contracts": trim_list(detail.get("contracts"), 8),
        "orders": trim_list(detail.get("orders"), 8),
        "invoices": trim_list(detail.get("invoices"), 8),
        "related_transactions": trim_list(detail.get("related_transactions"), 8),
        "guarantees": trim_list(detail.get("guarantees"), 8),
        "industry_profile": detail.get("industry_profile"),
        "peer_comparisons": trim_list(detail.get("peer_comparisons"), 8),
        "public_risks": trim_list(detail.get("public_risks"), 8),
        "validation_findings": trim_list(detail.get("validation_findings"), 8),
        "due_diligence_materials": [
            {
                "name": item["name"],
                "bucket_title": item["bucket_title"],
                "uploaded_at": item["uploaded_at"],
                "description": compact_text(item.get("description"), 300),
            }
            for item in trim_list(serialize_material_files(company_code), 12)
        ],
        "public_info_enrichment": public_info,
        "selected_knowledge_file_names": [item["name"] for item in knowledge_files],
        "knowledge_files": [
            {
                "id": item["id"],
                "name": item["name"],
                "category": item["category"],
                "owner": item["owner"],
                "description": item["description"],
                "content_text": compact_text(item.get("content_text", ""), 1200),
            }
            for item in trim_list(knowledge_files, 12)
        ],
        "selected_data_sources": summarize_selected_data_sources(selected_data_source_ids or []),
    }
    return {
        key: value
        for key, value in payload.items()
        if key in {"selected_data_sources", "knowledge_files", "selected_knowledge_file_names"} or key in selected
    }


def build_due_diligence_prompt(detail: dict, company_code: str, version_number: int, knowledge_files: list[dict], selected_data_source_ids: list[str] | None = None) -> str:
    payload = build_ai_input_payload(detail, company_code, knowledge_files, selected_data_source_ids)
    references = build_reference_catalog(detail, company_code, knowledge_files, selected_data_source_ids)
    sections = [
        "你是银行公司金融授信条线的高级尽职调查报告撰写智能体。",
        "",
        "你的任务是：",
        "基于输入的企业结构化数据、客户经理补充材料摘要、知识库引用文件摘要、规则核验结果，生成一份正式、专业、审慎、可用于银行内部授信审查流转的《尽职调查报告》。",
        "",
        "你的写作目标：",
        "1. 保持银行内部正式文风，语言客观、审慎、完整，不使用营销化表达。",
        "2. 先写事实，再写分析，再写判断，最后形成风险结论与建议。",
        "3. 所有结论必须尽量以输入数据为依据，不得凭空虚构。",
        "4. 缺失数据不要编造，统一写成“根据目前已获取资料，相关信息待进一步补充核实”。",
        "5. 如果不同数据之间存在矛盾、异常或勾稽不一致，必须明确写出并提示人工复核。",
        "6. 输出必须是完整长报告，不是提纲，不是摘要。",
        "7. 报告要贴近银行授信尽调场景，而不是债券承销场景。",
        "8. 知识库文件只作为写作口径、合规依据、风险提示参考，不能捏造其中不存在的结论。",
        "9. 只能参考本次明确选中的知识库文件；未选中的知识库文件不得自行假定或引用。",
        "10. 如果引用知识库内容，应以“根据《文件名》要求/口径”这样的方式自然融入，不要大段照抄。",
        "11. 对无法确认的事项，要使用审慎措辞，例如“需进一步核实”“建议补充提供”“目前未见充分证据支持”“需结合后续审查进一步判断”。",
        "12. 如果存在公开信息补全层，请充分吸收企业画像、年报链接摘要、行业洞察、治理摘要、公开风险和年度财务概览，扩充报告的背景描述与分析深度。",
        "13. 全文必须使用纯中文，不得输出英文章节名、英文表头、英文风险标签。",
        "14. 各章节必须展开为完整自然段，不能只写一句判断。",
        "15. 报告需要适度穿插 Markdown 表格，表格承载事实，正文承载判断。",
        "16. 关键事实、风险判断和补件建议必须在句末加入引用编号，例如“……需进一步核实[R1]”。",
        "17. 引用编号只能来自下方 reference_catalog，不得输出裸文件路径，不得自造不存在的引用编号。",
        "",
        "输出格式要求：",
        "- 使用 Markdown 输出。",
        "- 一级标题固定为：# 尽职调查报告。",
        "- 报告必须包含以下章节，并按此顺序输出。",
        "- 每章都要有完整正文，不能只列要点。",
        "- 每章必须固定包含四个子标题：### 风险定性、### 核查证据、### 财务与合规影响评估、### 补充尽调建议。",
        "- 最后一章必须输出明确的尽调结论与建议方向。",
        "",
        "报告章节结构如下：",
        "# 尽职调查报告",
        "## 第一章 基本情况与历史沿革核查",
        "## 第二章 主营业务与经营模式深挖",
        "## 第三章 财务与会计主要数据指标评析",
        "## 第四章 资产结构穿透与资金独立性分析",
        "## 第五章 有息债务期限结构与表外负债排查",
        "## 第六章 初步尽调结论与补件建议",
        "",
        "章节展开要求：",
        "- 每个章节必须有完整自然段，不允许只写一句话。",
        "- 对上市公司或公开资料较丰富的客户，应充分利用公开信息补全层补写背景、治理、行业和公开风险信息。",
        "- 在第三章至少插入 2 张表格，优先包含核心财务指标表、收入回款税票勾稽表。",
        "- 在第五章至少插入 1 张表格，优先包含债务结构表或担保/或有事项表。",
        "- 全文至少输出 3 张 Markdown 表格。",
        "- 缺资料时不得略过章节，必须写明当前缺口及其对判断的影响。",
        "- 在第六章必须输出综合判断、建议补充材料清单，以及“建议进入下一步审查 / 建议补件后再审 / 现阶段暂不建议推进”三选一结论。",
        "",
        "写作要求补充：",
        "1. 不得输出“模型认为”“AI判断”等措辞。",
        "2. 不要写空洞套话，要尽量结合输入数据。",
        "3. 对没有数据支撑的章节，简洁写明“资料待进一步补充核实”，但如果公开信息补全层中有相关内容，应优先吸收后再写。",
        "4. 如果知识库文件与企业情况有关联，可自然引用，例如“根据《对公授信准入政策（2026版）》相关口径……”。",
        "5. 如果规则核验命中异常，必须在相应章节中呼应，并同时写清事实、影响和建议。",
        "6. 全文结论要前后一致，不能前面写风险高、后面又直接建议通过。",
        "7. 输出必须是完整正文，不要输出 JSON，不要输出解释说明，不要输出“以下是报告”。",
        "8. 表格表头必须使用中文，不得出现 Metric、Revenue、Change 等英文表头。",
        "9. 使用引用时统一写成 [R1]、[R2]，同一句可使用多个编号，例如 [R2][R5]。",
        "",
        "后续预审视角补充：",
        "生成正文时，请尽量让下列判断点在报告中有事实依据、判断过程或待补数据提示，便于后续预审同步完成。",
        *[
            f"{index + 1}. {item['judgement_focus']} 需要数据：{item['required_data']} 完善方向：{item['improvement_direction']}"
            for index, item in enumerate(REVIEW_JUDGEMENT_GUIDANCE)
        ],
        "",
        f"当前生成版本：V{version_number}",
        "",
        "本次已明确选中的知识库文件如下，写作时只能参考这些文件：",
        "；".join(payload.get("selected_knowledge_file_names") or ["本次未选择知识库文件"]),
        "",
        "请严格参考以下 Markdown 模板骨架组织正文：",
        REPORT_MARKDOWN_TEMPLATE,
        "",
        "下面是本次生成报告的输入数据，请严格基于这些数据写作：",
        "",
        reference_catalog_prompt_block(references),
        json_block("selected_data_sources", payload.get("selected_data_sources")),
        json_block("company", payload.get("company")),
        json_block("case_info", payload.get("case_info")),
        json_block("recommendation", payload.get("recommendation")),
        json_block("people", payload.get("people")),
        json_block("related_companies", payload.get("related_companies")),
        json_block("shareholding_changes", payload.get("shareholding_changes")),
        json_block("profile_attributes", payload.get("profile_attributes")),
        json_block("latest_metrics", payload.get("latest_metrics")),
        json_block("recent_metric_history", payload.get("recent_metric_history")),
        json_block("reconciliation_checks", payload.get("reconciliation_checks")),
        json_block("tax_invoice_checks", payload.get("tax_invoice_checks")),
        json_block("tax_filings", payload.get("tax_filings")),
        json_block("bank_summaries", payload.get("bank_summaries")),
        json_block("credit_history", payload.get("credit_history")),
        json_block("contracts", payload.get("contracts")),
        json_block("orders", payload.get("orders")),
        json_block("invoices", payload.get("invoices")),
        json_block("related_transactions", payload.get("related_transactions")),
        json_block("guarantees", payload.get("guarantees")),
        json_block("industry_profile", payload.get("industry_profile")),
        json_block("peer_comparisons", payload.get("peer_comparisons")),
        json_block("public_risks", payload.get("public_risks")),
        json_block("validation_findings", payload.get("validation_findings")),
        json_block("due_diligence_materials", payload.get("due_diligence_materials")),
        json_block("public_info_enrichment", payload.get("public_info_enrichment")),
        json_block("selected_knowledge_file_names", payload.get("selected_knowledge_file_names")),
        json_block("knowledge_files", payload.get("knowledge_files")),
        "",
        "如果某一字段为空，请按“资料待进一步补充核实”处理，不得编造。",
        "现在开始直接输出完整《尽职调查报告》正文。",
    ]
    return "\n".join(sections)


def build_skill_pipeline_report_prompt(
    detail: dict,
    company_code: str,
    version_number: int,
    knowledge_files: list[dict],
    selected_data_source_ids: list[str] | None = None,
) -> str:
    payload = build_ai_input_payload(detail, company_code, knowledge_files, selected_data_source_ids)
    audit_summary = build_skill_audit_summary(company_code, detail)
    references = build_reference_catalog(detail, company_code, knowledge_files, selected_data_source_ids)
    auditing_skill = load_skill_file("auditing-data")
    writing_skill = load_skill_file("writing-report")
    template_text = load_skill_file("writing-report", "REPORT_TEMPLATE.md") or REPORT_MARKDOWN_TEMPLATE

    sections = [
        "请根据以下三份说明共同完成报告生成：",
        "1. auditing-data：定义审计阶段输出了哪些 findings，以及这些 findings 应如何被吸收。",
        "2. writing-report：定义了尽调报告的文风、章节、表格和锚点规则。",
        "3. REPORT_TEMPLATE.md：定义了最终报告的 Markdown 结构骨架。",
        "",
        "写作硬约束：",
        "1. 必须输出纯中文 Markdown。",
        "2. 必须完整输出 6 章，且每章包含“风险定性 / 核查证据 / 财务与合规影响评估 / 补充尽调建议”。",
        "3. 必须穿插表格，尤其在财务勾稽、关联关系、债务担保、补件建议处。",
        "4. 必须优先吸收审计结果中的高风险 findings。",
        "5. 关键风险判断后要自然附上 reference_catalog 中的引用编号。",
        "6. 缺资料时不能跳过章节，必须写明资料缺口和建议补件。",
        "7. 不得输出英文章节名、英文表头、英文风险标题。",
        "8. 只能参考本次已明确选中的知识库文件；未选中的知识库文件不得自行引用。",
        "9. 最终正文里的来源锚点必须统一输出为 reference_catalog 中的 [R1]、[R2] 编号，不得输出裸文件路径。",
        "10. 对收入真实性、关联交易、担保、债务、资产瑕疵等关键判断，句末必须附至少一个 [R] 编号。",
        "",
        f"当前公司：{detail.get('company', {}).get('name') or company_code}",
        f"当前生成版本：V{version_number}",
        "本次已明确选中的知识库文件：",
        "；".join(payload.get("selected_knowledge_file_names") or ["本次未选择知识库文件"]),
        "",
        "以下为审计技能说明：",
        auditing_skill,
        "",
        "以下为写稿技能说明：",
        writing_skill,
        "",
        "以下为报告模板：",
        template_text,
        "",
        "以下为本轮审计结果（写稿时必须优先吸收）：",
        json_block("audit_summary", audit_summary),
        "",
        "以下为本轮可使用的引用目录，正文只能使用这些编号：",
        reference_catalog_prompt_block(references),
        "",
        "以下为本轮写稿可使用的业务输入：",
        json_block("selected_data_sources", payload.get("selected_data_sources")),
        json_block("company", payload.get("company")),
        json_block("case_info", payload.get("case_info")),
        json_block("recommendation", payload.get("recommendation")),
        json_block("people", payload.get("people")),
        json_block("related_companies", payload.get("related_companies")),
        json_block("shareholding_changes", payload.get("shareholding_changes")),
        json_block("profile_attributes", payload.get("profile_attributes")),
        json_block("latest_metrics", payload.get("latest_metrics")),
        json_block("recent_metric_history", payload.get("recent_metric_history")),
        json_block("reconciliation_checks", payload.get("reconciliation_checks")),
        json_block("tax_invoice_checks", payload.get("tax_invoice_checks")),
        json_block("tax_filings", payload.get("tax_filings")),
        json_block("bank_summaries", payload.get("bank_summaries")),
        json_block("credit_history", payload.get("credit_history")),
        json_block("contracts", payload.get("contracts")),
        json_block("orders", payload.get("orders")),
        json_block("invoices", payload.get("invoices")),
        json_block("related_transactions", payload.get("related_transactions")),
        json_block("guarantees", payload.get("guarantees")),
        json_block("industry_profile", payload.get("industry_profile")),
        json_block("peer_comparisons", payload.get("peer_comparisons")),
        json_block("public_risks", payload.get("public_risks")),
        json_block("validation_findings", payload.get("validation_findings")),
        json_block("due_diligence_materials", payload.get("due_diligence_materials")),
        json_block("public_info_enrichment", payload.get("public_info_enrichment")),
        json_block("selected_knowledge_file_names", payload.get("selected_knowledge_file_names")),
        json_block("knowledge_files", payload.get("knowledge_files")),
        "",
        "如果某一字段为空，请按“资料待进一步补充核实”处理，不得编造。",
        "现在开始直接输出完整《尽职调查报告》正文。",
    ]
    return "\n".join(sections)


def parse_markdown_report(markdown_text: str) -> tuple[list[dict], str]:
    normalized = markdown_text.replace("\r", "").strip()
    lines = normalized.split("\n")
    sections: list[dict] = []
    current_title = None
    current_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if line.startswith("## "):
            if current_title:
                sections.append(
                    {
                        "id": f"section-{len(sections) + 1}",
                        "title": current_title,
                        "content": "\n".join(current_lines).strip(),
                    }
                )
            current_title = line.replace("## ", "", 1).strip()
            current_lines = []
            continue
        if line.startswith("# "):
            continue
        current_lines.append(raw_line)

    if current_title:
        sections.append(
            {
                "id": f"section-{len(sections) + 1}",
                "title": current_title,
                "content": "\n".join(current_lines).strip(),
            }
        )

    full_text = "\n\n".join(
        [f"{section['title']}\n{section['content']}".strip() for section in sections]
    )
    return sections, full_text


def parse_editor_report_text(text: str) -> tuple[list[dict], str]:
    normalized = (text or "").replace("\r", "").strip()
    if "## " in normalized:
        sections, full_text = parse_markdown_report(normalized)
        if sections:
            return sections, full_text
    return (
        [
            {
                "id": "editor-section-1",
                "title": "尽调报告正文",
                "content": normalized or "当前版本暂无正文。",
            }
        ],
        normalized or "当前版本暂无正文。",
    )


def parse_json_response_text(text: str) -> dict:
    normalized = (text or "").strip()
    if normalized.startswith("```"):
        normalized = normalized.strip("`")
        if normalized.startswith("json"):
            normalized = normalized[4:].strip()
    start = normalized.find("{")
    end = normalized.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise RuntimeError("DeepSeek 预审结果未返回有效 JSON")
    return json.loads(normalized[start : end + 1])


def generate_report_with_deepseek(detail: dict, company_code: str, version_number: int, knowledge_files: list[dict], selected_data_source_ids: list[str] | None = None) -> tuple[list[dict], str]:
    prompt = build_due_diligence_prompt(detail, company_code, version_number, knowledge_files, selected_data_source_ids)
    references = build_reference_catalog(detail, company_code, knowledge_files, selected_data_source_ids)
    response = call_deepseek_chat(
        [{"role": "user", "content": prompt}],
        system_prompt=REPORT_SYSTEM_PROMPT,
        model=DEEPSEEK_MODEL,
    )
    content = (
        response.get("choices", [{}])[0]
        .get("message", {})
        .get("content", "")
        .strip()
    )
    if not content:
        raise RuntimeError("DeepSeek 未返回报告正文")
    content = normalize_reference_markers(content, references)
    sections, full_text = parse_markdown_report(content)
    if not sections:
        raise RuntimeError("DeepSeek 返回内容无法解析为报告章节")
    return sections, full_text


def build_review_prompt(detail: dict, company_code: str, version: dict, report_text: str, knowledge_files: list[dict]) -> str:
    payload = build_ai_input_payload(detail, company_code, knowledge_files, version.get("selected_data_source_ids"))
    sections = [
        "你是银行授信尽调预审智能体，请基于当前尽调报告正文、结构化企业数据、已选外部法规、已选内部制度、已选专家经验，输出一份预审风险提示清单。",
        "",
        "你的任务要求：",
        "1. 必须识别当前报告中需要补充判断、需要补充数据、需要修订表述的地方。",
        "2. 已经提供的外部法规、内部制度、专家经验文件，视为“已选依据”，不得错误输出成“未上传”“未提供”。",
        "3. 风险提示必须贴近银行客户经理实际工作，语言明确、简洁、可执行。",
        "4. 优先关注：主体准入、治理结构、经营真实性、还款来源、财务质量、行业政策、关联交易、担保缓释、合规边界、资料缺口、结论一致性。",
        "5. 不要重复同一问题；每条提示应聚焦一个核心判断点。",
        "6. 如果当前依据已经足够，不要机械要求补材料；应结合报告内容给出更细的判断或修订建议。",
        "7. 输出必须是 JSON，不要输出 Markdown，不要输出解释说明。",
        "",
        "输出 JSON 结构如下：",
        "{",
        '  "overall_result": "一句话总结，说明当前报告预审状态",',
        '  "review_summary": "2-4句中文，概括当前报告主要优点、主要缺口和建议动作",',
        '  "findings": [',
        "    {",
        '      "severity": "high|medium|low",',
        '      "title": "短标题",',
        '      "issue": "当前问题描述",',
        '      "judgement_focus": "客户经理需要做出的判断",',
        '      "required_data": "若还需数据，写清楚；若当前资料基本够用，可写“现有资料可支持初判，建议补充xx增强结论”",',
        '      "improvement_direction": "建议如何修改报告或补充分析",',
        '      "source_type": "外部法规|内部制度|专家经验|结构化数据|综合判断",',
        '      "source_basis": "命中的依据文件名或数据点" ',
        "    }",
        "  ]",
        "}",
        "",
        "findings 数量要求：输出 8 到 15 条，按重要性排序。",
        "如果某类问题已被充分覆盖，可以不强行凑数，但不要少于 8 条。",
        "",
        "下面是输入数据：",
        json_block("current_report_text", report_text),
        json_block("selected_knowledge_files", payload.get("knowledge_files")),
        json_block("company", payload.get("company")),
        json_block("case_info", payload.get("case_info")),
        json_block("recommendation", payload.get("recommendation")),
        json_block("people", payload.get("people")),
        json_block("related_companies", payload.get("related_companies")),
        json_block("shareholding_changes", payload.get("shareholding_changes")),
        json_block("latest_metrics", payload.get("latest_metrics")),
        json_block("recent_metric_history", payload.get("recent_metric_history")),
        json_block("reconciliation_checks", payload.get("reconciliation_checks")),
        json_block("tax_invoice_checks", payload.get("tax_invoice_checks")),
        json_block("bank_summaries", payload.get("bank_summaries")),
        json_block("credit_history", payload.get("credit_history")),
        json_block("contracts", payload.get("contracts")),
        json_block("orders", payload.get("orders")),
        json_block("invoices", payload.get("invoices")),
        json_block("related_transactions", payload.get("related_transactions")),
        json_block("guarantees", payload.get("guarantees")),
        json_block("industry_profile", payload.get("industry_profile")),
        json_block("peer_comparisons", payload.get("peer_comparisons")),
        json_block("public_risks", payload.get("public_risks")),
        json_block("validation_findings", payload.get("validation_findings")),
        json_block("due_diligence_materials", payload.get("due_diligence_materials")),
        json_block("public_info_enrichment", payload.get("public_info_enrichment")),
        "",
        "现在开始，只输出 JSON。",
    ]
    return "\n".join(sections)


def build_ai_review_result(detail: dict, company_code: str, version: dict, report_text: str, knowledge_files: list[dict]) -> dict:
    prompt = build_review_prompt(detail, company_code, version, report_text, knowledge_files)
    response = call_deepseek_chat(
        [{"role": "user", "content": prompt}],
        system_prompt=REVIEW_SYSTEM_PROMPT,
        model=DEEPSEEK_MODEL,
    )
    content = (
        response.get("choices", [{}])[0]
        .get("message", {})
        .get("content", "")
        .strip()
    )
    if not content:
        raise RuntimeError("DeepSeek 未返回预审结果")
    payload = parse_json_response_text(content)
    findings = payload.get("findings") or []
    normalized_findings = []
    for index, item in enumerate(findings[:15]):
        severity = str(item.get("severity") or "medium").lower()
        if severity not in {"high", "medium", "low"}:
            severity = "medium"
        normalized_findings.append(
            {
                "severity": severity,
                "source_rule": compact_text(item.get("source_basis"), 240),
                "source_type": compact_text(item.get("source_type"), 60),
                "section_id": f"ai-review-{index + 1}",
                "section_title": compact_text(item.get("title"), 60),
                "issue": compact_text(item.get("issue"), 400),
                "suggestion": compact_text(item.get("improvement_direction"), 400),
                "judgement_focus": compact_text(item.get("judgement_focus"), 240),
                "required_data": compact_text(item.get("required_data"), 240),
                "improvement_direction": compact_text(item.get("improvement_direction"), 400),
            }
        )
    if not normalized_findings:
        raise RuntimeError("DeepSeek 预审结果缺少 findings")
    counts = {
        "high": sum(1 for item in normalized_findings if item["severity"] == "high"),
        "medium": sum(1 for item in normalized_findings if item["severity"] == "medium"),
        "low": sum(1 for item in normalized_findings if item["severity"] == "low"),
    }
    return {
        "review_status": "已预审",
        "overall_result": compact_text(payload.get("overall_result"), 120),
        "reviewed_at": now_iso(),
        "counts": counts,
        "review_summary": compact_text(payload.get("review_summary"), 400),
        "findings": normalized_findings,
    }


def build_generation_basis(detail: dict, knowledge_files: list[dict], company_code: str, selected_data_source_ids: list[str] | None = None) -> list[str]:
    materials = serialize_material_files(company_code)
    public_info = load_public_enrichment(company_code)
    public_blocks = sum(1 for value in public_info.values() if isinstance(value, str) and value.strip())
    selected_groups = summarize_selected_data_sources(selected_data_source_ids or [])
    basis = [
        f"客户基础信息、治理、经营与财务数据已从数据库同步，共 {len(detail.get('latest_metrics') or [])} 项核心财务指标",
        f"真实性与核验结果已纳入，共 {len(detail.get('validation_findings') or [])} 条核验发现",
        f"尽调补充资料已纳入，共 {len(materials)} 份分组资料",
        f"公开信息补全层已纳入，共 {public_blocks} 组公开资料摘要",
        f"知识库引用文件共 {len(knowledge_files)} 份",
        f"本次关联结构化数据库内容共 {sum(len(group['items']) for group in selected_groups)} 项",
    ]
    if knowledge_files:
        basis.append("本次引用知识文件：" + "、".join(item["name"] for item in knowledge_files))
    if selected_groups:
        basis.append(
            "本次关联数据分组："
            + "；".join(f"{group['group_title']}（{'、'.join(group['items'])}）" for group in selected_groups)
        )
    return basis


def build_skill_generation_basis(detail: dict, company_code: str, knowledge_files: list[dict], selected_data_source_ids: list[str] | None = None) -> list[str]:
    binding = get_skill_pipeline_binding(company_code)
    base_basis = build_generation_basis(detail, knowledge_files, company_code, selected_data_source_ids)
    if not binding:
        return base_basis
    skill_basis = [
        f"已调用 Skill Pipeline：{binding['skills'][0]} -> {binding['skills'][1]}",
        "审计阶段先输出结构化 findings，再由 DeepSeek 按 skill 说明书生成风险前置长文本报告" if get_deepseek_api_key() else "当前未检测到 DeepSeek Key，已退回本地模板生成",
    ]
    if binding.get("mock_vdr_dir"):
        skill_basis.append(f"Mock VDR 目录：{binding['mock_vdr_dir']}")
    return skill_basis + base_basis


def build_review_result(detail: dict, version: dict) -> dict:
    files = version["review_files"]
    internal_files = [item for item in files if item["category"] == "internal"]
    external_files = [item for item in files if item["category"] == "external"]
    findings: list[dict] = []
    knowledge_map = knowledge_file_map()
    expert_files = [
        knowledge_map[file_id]["name"]
        for file_id in version.get("knowledge_file_ids", [])
        if file_id in knowledge_map and knowledge_map[file_id].get("category") == "experience"
    ]

    for item in REVIEW_JUDGEMENT_GUIDANCE:
        improvement = item["improvement_direction"]
        if item["source_type"] == "内外规符合性":
            if not internal_files:
                improvement += " 当前尚未上传银行内规文件。"
            if not external_files:
                improvement += " 当前尚未上传外规文件。"
        if item["source_type"] == "行业外部环境" and not expert_files:
            improvement += " 当前版本尚未关联专家经验文件，可补充经验库口径。"
        findings.append(
            {
                "severity": item["severity"],
                "source_rule": item["required_data"],
                "source_type": item["source_type"],
                "section_id": item["id"],
                "section_title": item["source_type"],
                "issue": item["judgement_focus"],
                "suggestion": improvement,
                "judgement_focus": item["judgement_focus"],
                "required_data": item["required_data"],
                "improvement_direction": improvement,
            }
        )

    counts = {
        "high": sum(1 for item in findings if item["severity"] == "high"),
        "medium": sum(1 for item in findings if item["severity"] == "medium"),
        "low": sum(1 for item in findings if item["severity"] == "low"),
    }
    overall_result = "高风险待复核" if counts["high"] else "存在待完善项"
    return {
        "review_status": "已预审",
        "overall_result": overall_result,
        "reviewed_at": now_iso(),
        "counts": counts,
        "findings": findings,
    }


def generate_report_with_skill_pipeline(
    company_code: str,
    version_number: int,
    detail: dict | None = None,
    knowledge_files: list[dict] | None = None,
    selected_data_source_ids: list[str] | None = None,
) -> tuple[list[dict], str]:
    if not get_skill_pipeline_binding(company_code):
        raise RuntimeError(f"No skill pipeline configured for {company_code}")
    detail = detail or get_company_detail_base_payload(company_code)
    if detail is None:
        raise RuntimeError(f"Company detail not found for {company_code}")

    if not get_deepseek_api_key() and has_antigravity_mock_pack(company_code):
        markdown = build_antigravity_report_markdown(company_code, version_number, detail)
        references = build_reference_catalog(detail, company_code, knowledge_files or [], selected_data_source_ids)
        markdown = normalize_reference_markers(markdown, references)
        sections, full_text = parse_markdown_report(markdown)
        if not sections:
            raise RuntimeError("Skill pipeline fallback returned no report sections")
        return sections, full_text
    if not get_deepseek_api_key():
        sections, full_text, _ = build_fallback_report_sections(detail, version_number)
        if not sections:
            raise RuntimeError("Skill pipeline template fallback returned no report sections")
        return sections, full_text

    prompt = build_skill_pipeline_report_prompt(
        detail,
        company_code,
        version_number,
        knowledge_files or [],
        selected_data_source_ids,
    )
    response = call_deepseek_chat(
        [{"role": "user", "content": prompt}],
        system_prompt=SKILL_PIPELINE_REPORT_SYSTEM_PROMPT,
        model=DEEPSEEK_MODEL,
    )
    markdown = (
        response.get("choices", [{}])[0]
        .get("message", {})
        .get("content", "")
        .strip()
    )
    references = build_reference_catalog(detail, company_code, knowledge_files or [], selected_data_source_ids)
    markdown = normalize_reference_markers(markdown, references)
    sections, full_text = parse_markdown_report(markdown)
    if not sections:
        raise RuntimeError("Skill pipeline DeepSeek output could not be parsed into report sections")
    return sections, full_text


def serialize_material_files(company_code: str) -> list[dict]:
    runtime = ensure_material_runtime(company_code)
    return runtime["files"]


def ensure_report_runtime(company_code: str, detail: dict) -> dict:
    runtime = REPORT_RUNTIME.get(company_code)
    if runtime:
        return runtime

    if get_skill_pipeline_binding(company_code):
        references = build_reference_catalog(detail, company_code, [], get_all_report_data_source_ids())
        sections, full_text = generate_report_with_skill_pipeline(
            company_code,
            1,
            detail,
            [],
            get_all_report_data_source_ids(),
        )
        sections = normalize_report_section_references(sections, references)
        full_text = normalize_reference_markers(full_text, references)
        based_on = build_skill_generation_basis(detail, company_code, [], get_all_report_data_source_ids())
        generation_mode = "skill_pipeline"
    else:
        sections, full_text, based_on = build_fallback_report_sections(detail, 1)
        references = build_reference_catalog(detail, company_code, [], get_all_report_data_source_ids())
        generation_mode = "template"
    runtime = {
        "next_version": 2,
        "versions": [
            {
                "id": "v1",
                "version_label": "V1",
                "created_at": now_iso(),
                "status": "已生成",
                "based_on": based_on,
                "section_list": sections,
                "full_text": full_text,
                "knowledge_file_ids": [],
                "selected_data_source_ids": get_all_report_data_source_ids(),
                "references": references,
                "review_files": [],
                "review_result": None,
                "review_edit_text": full_text,
                "review_saved_at": None,
                "generation_mode": generation_mode,
            }
        ],
    }
    REPORT_RUNTIME[company_code] = runtime
    return runtime


def get_companies_payload() -> dict:
    workbook_index = get_workbook_index()
    with get_connection() as conn:
        companies = fetch_all(
            conn,
            """
            SELECT
              c.company_code,
              c.name,
              c.region_province,
              c.region_city,
              c.industry_category,
              c.subindustry,
              c.enterprise_scale,
              c.operating_status,
              c.risk_tier,
              c.overview,
              COUNT(DISTINCT d.id) AS case_count,
              COUNT(DISTINCT r.person_id) AS people_count,
              COUNT(DISTINCT vf.id) AS finding_count,
              cr.recommendation_status,
              cr.note AS recommendation_note
            FROM companies c
            LEFT JOIN due_diligence_cases d ON d.company_id = c.id
            LEFT JOIN company_person_roles r ON r.company_id = c.id
            LEFT JOIN validation_findings vf ON vf.company_id = c.id
            LEFT JOIN credit_recommendations cr ON cr.company_id = c.id
            GROUP BY c.id
            ORDER BY c.company_code
            """
        )

    for item in companies:
        workbook_info = workbook_index["companies"].get(item["company_code"], {})
        item["workbook_sheet_count"] = len(workbook_info.get("sheet_counts", {}))
        item["data_mode"] = "skill_pipeline" if get_skill_pipeline_binding(item["company_code"]) else (
            "public_pack" if item.get("recommendation_status") == "public_pack_only" else "full_case"
        )
    return {
        "sqlite_path": str(DB_PATH),
        "excel_path": workbook_index["path"],
        "sheet_names": workbook_index["sheet_names"],
        "companies": companies,
    }


def build_home_feed_payload() -> dict:
    stage_labels = {
        "report_review": "洞察复核中",
        "committee_ready": "洞察已完成",
        "supplement_pending": "待补充资料",
        "data_pack_ready": "数据包就绪",
    }
    severity_labels = {
        "high": ("高风险", "status-pressure"),
        "medium": ("中风险", "status-watch"),
        "low": ("低风险", "status-blue"),
    }

    with get_connection() as conn:
        recent_investigations = fetch_all(
            conn,
            """
            SELECT
              c.company_code,
              c.name,
              COALESCE(c.subindustry, c.industry_category, '所属行业') AS industry,
              d.product_type,
              d.current_stage,
              d.application_date
            FROM due_diligence_cases d
            JOIN companies c ON c.id = d.company_id
            WHERE d.product_type != 'listed_company_public_pack'
            ORDER BY d.application_date DESC, d.created_at DESC
            LIMIT 3
            """,
        )

        risk_alerts = fetch_all(
            conn,
            """
            SELECT
              c.company_code,
              c.name,
              e.severity,
              e.title,
              e.summary,
              e.event_date
            FROM public_risk_events e
            JOIN companies c ON c.id = e.company_id
            ORDER BY
              CASE e.severity
                WHEN 'high' THEN 3
                WHEN 'medium' THEN 2
                ELSE 1
              END DESC,
              e.event_date DESC,
              e.id DESC
            LIMIT 3
            """,
        )

    investigations_payload = [
        {
            "company": item["name"],
            "tag": "风险洞察" if item.get("product_type") == "micro_enterprise_credit" else "企业洞察",
            "industry": item.get("industry") or "所属行业",
            "status": stage_labels.get(item.get("current_stage"), item.get("current_stage") or "进行中"),
            "time": item.get("application_date") or "—",
        }
        for item in recent_investigations
    ]

    alerts_payload = []
    for item in risk_alerts:
        level, level_class = severity_labels.get(item.get("severity"), ("风险关注", "status-blue"))
        alerts_payload.append(
            {
                "company": item["name"],
                "level": level,
                "levelClass": level_class,
                "detail": item.get("summary") or item.get("title") or "请进入企业详情查看最新风险动态。",
                "time": item.get("event_date") or "—",
            }
        )

    return {
        "recentInvestigations": investigations_payload,
        "riskAlerts": alerts_payload,
    }


def get_company_detail_base_payload(company_code: str) -> dict | None:
    workbook_index = get_workbook_index()
    workbook_company = workbook_index["companies"].get(company_code, {"sheet_counts": {}, "samples": {}})

    with get_connection() as conn:
        company = fetch_one(
            conn,
            """
            SELECT
              id,
              company_code,
              name,
              unified_social_credit_code,
              established_on,
              region_province,
              region_city,
              industry_category,
              subindustry,
              enterprise_scale,
              operating_status,
              registered_capital_cny,
              paid_in_capital_cny,
              overview,
              risk_tier
            FROM companies
            WHERE company_code = ?
            """,
            (company_code,),
        )
        if not company:
            return None

        company_id = company["id"]
        case_info = fetch_one(
            conn,
            """
            SELECT
              d.case_no,
              d.case_name,
              d.application_date,
              d.product_type,
              d.requested_amount_cny,
              d.requested_term_months,
              d.current_stage,
              d.decision_status,
              la.product_name,
              la.approved_amount_cny,
              la.loan_purpose,
              la.primary_repayment_source,
              la.secondary_repayment_source,
              la.repayment_method,
              la.guarantee_mode,
              la.annual_rate_min,
              la.annual_rate_max,
              la.tenor_months
            FROM due_diligence_cases d
            LEFT JOIN loan_applications la ON la.case_id = d.id
            WHERE d.company_id = ?
            ORDER BY d.application_date DESC
            LIMIT 1
            """,
            (company_id,),
        )

        recommendation = fetch_one(
            conn,
            """
            SELECT
              recommendation_status,
              suggested_amount_cny,
              suggested_term_months,
              suggested_rate_min,
              suggested_rate_max,
              guarantee_requirement,
              supplemental_requirements,
              rejection_reason,
              note
            FROM credit_recommendations
            WHERE company_id = ?
            LIMIT 1
            """,
            (company_id,),
        )

        people = fetch_all(
            conn,
            """
            SELECT
              p.full_name,
              r.role_type,
              r.title,
              r.equity_ratio,
              r.voting_ratio,
              r.is_actual_controller,
              r.is_guarantor,
              p.education_level,
              p.birth_year
            FROM company_person_roles r
            JOIN persons p ON p.id = r.person_id
            WHERE r.company_id = ?
            ORDER BY r.is_actual_controller DESC, COALESCE(r.equity_ratio, 0) DESC, p.full_name
            """,
            (company_id,),
        )

        related_companies = fetch_all(
            conn,
            """
            SELECT related_company_name, relation_type, risk_flag, control_path, note
            FROM related_companies
            WHERE company_id = ?
            ORDER BY risk_flag DESC, related_company_name
            LIMIT 8
            """,
            (company_id,),
        )

        shareholding_changes = fetch_all(
            conn,
            """
            SELECT change_date, change_type, declared_reason, risk_comment
            FROM shareholding_changes
            WHERE company_id = ?
            ORDER BY change_date DESC
            LIMIT 6
            """,
            (company_id,),
        )

        profile_attributes = fetch_all(
            conn,
            """
            SELECT attribute_group, label, value_text, source_ref, note
            FROM company_profile_attributes
            WHERE company_id = ?
            ORDER BY attribute_group, label
            LIMIT 40
            """,
            (company_id,),
        )

        findings = fetch_all(
            conn,
            """
            SELECT severity, finding_title, finding_summary, impact_summary, confidence, requires_manual_review
            FROM validation_findings
            WHERE company_id = ?
            ORDER BY CASE severity WHEN 'high' THEN 3 WHEN 'medium' THEN 2 ELSE 1 END DESC, id
            LIMIT 12
            """,
            (company_id,),
        )

        report_sections = fetch_all(
            conn,
            """
            SELECT section_code, section_title, content, confidence
            FROM report_sections
            WHERE company_id = ?
            ORDER BY display_order
            """,
            (company_id,),
        )

        latest_financial_period = fetch_one(
            conn,
            """
            SELECT id, period_code
            FROM financial_periods
            WHERE company_id = ?
            ORDER BY CASE WHEN period_code LIKE '%FY' OR period_code LIKE '%1231' THEN 1 ELSE 0 END DESC, end_date DESC
            LIMIT 1
            """,
            (company_id,),
        )

        latest_metrics = []
        if latest_financial_period:
            latest_metrics = fetch_all(
                conn,
                """
                SELECT metric_code, metric_name, metric_category, value, unit
                FROM financial_metrics
                WHERE company_id = ? AND period_id = ?
                  AND metric_code IN (
                    'revenue','gross_profit','gross_margin_pct','net_profit','cash','operating_cash_flow',
                    'operating_cash_flow_margin_pct','accounts_receivable','accounts_payable','inventory',
                    'total_assets','total_liabilities','net_assets','asset_liability_ratio',
                    'capital_expenditure','fixed_assets','operating_cashflow','debt_to_assets_pct',
                    'shareholder_equity','operating_cost'
                  )
                ORDER BY metric_code
                """,
                (company_id, latest_financial_period["id"]),
            )

        recent_metric_history = fetch_all(
            conn,
            """
            SELECT fp.period_code, fm.metric_code, fm.metric_name, fm.value, fm.unit
            FROM financial_metrics fm
            JOIN financial_periods fp ON fp.id = fm.period_id
            WHERE fm.company_id = ?
              AND fm.metric_code IN ('revenue', 'net_profit', 'cash', 'accounts_receivable')
            ORDER BY fp.end_date DESC
            LIMIT 60
            """,
            (company_id,),
        )

        reconciliation_checks = fetch_all(
            conn,
            """
            SELECT check_name, lhs_label, lhs_value, rhs_label, rhs_value, variance_ratio, status, interpretation
            FROM financial_reconciliation_checks
            WHERE company_id = ?
            ORDER BY period_id DESC, id
            LIMIT 8
            """,
            (company_id,),
        )

        tax_invoice_checks = fetch_all(
            conn,
            """
            SELECT declared_revenue_cny, invoiced_amount_cny, bank_receipts_cny, gap_ratio, status, check_note
            FROM tax_invoice_consistency_checks
            WHERE company_id = ?
            ORDER BY period_id DESC, id
            LIMIT 6
            """,
            (company_id,),
        )

        tax_filings = fetch_all(
            conn,
            """
            SELECT filing_type, declared_revenue_cny, output_tax_cny, input_tax_cny, tax_burden_ratio, filing_status
            FROM tax_filings
            WHERE company_id = ?
            ORDER BY period_id DESC, id
            LIMIT 6
            """,
            (company_id,),
        )

        bank_summaries = fetch_all(
            conn,
            """
            SELECT inflow_total_cny, outflow_total_cny, net_flow_cny, average_daily_balance_cny,
                   top_inflow_counterparty, top_inflow_amount_cny, top_outflow_counterparty,
                   top_outflow_amount_cny, transaction_count, large_transaction_count, summary_note
            FROM bank_settlement_summaries
            WHERE company_id = ?
            ORDER BY period_id DESC, id
            LIMIT 6
            """,
            (company_id,),
        )

        receivables = fetch_all(
            conn,
            """
            SELECT record_type, balance_cny, overdue_over_90d_cny, top5_ratio_pct, average_days
            FROM receivables_payables
            WHERE company_id = ?
            ORDER BY period_id DESC, id
            LIMIT 10
            """,
            (company_id,),
        )

        social_security = fetch_all(
            conn,
            """
            SELECT insured_headcount, month_change, average_base_cny
            FROM social_security_metrics
            WHERE company_id = ?
            ORDER BY period_id DESC
            LIMIT 6
            """,
            (company_id,),
        )

        utilities = fetch_all(
            conn,
            """
            SELECT utility_type, consumption_value, unit, yoy_change_pct
            FROM utility_metrics
            WHERE company_id = ?
            ORDER BY period_id DESC, id
            LIMIT 8
            """,
            (company_id,),
        )

        logistics = fetch_all(
            conn,
            """
            SELECT shipment_count, shipment_weight_kg, freight_cost_cny, return_rate_pct
            FROM logistics_metrics
            WHERE company_id = ?
            ORDER BY period_id DESC
            LIMIT 6
            """,
            (company_id,),
        )

        guarantees = fetch_all(
            conn,
            """
            SELECT guarantee_type, asset_name, asset_category, appraised_value_cny, pledge_rate, lien_status, guarantee_status, notes
            FROM guarantees
            WHERE case_id IN (SELECT id FROM due_diligence_cases WHERE company_id = ?)
            ORDER BY id
            LIMIT 8
            """,
            (company_id,),
        )

        contracts = fetch_all(
            conn,
            """
            SELECT contract_no, contract_type, sign_date, amount_cny, payment_terms, performance_status
            FROM contracts
            WHERE company_id = ?
            ORDER BY sign_date DESC
            LIMIT 8
            """,
            (company_id,),
        )

        orders = fetch_all(
            conn,
            """
            SELECT order_no, order_date, order_amount_cny, delivery_due_date, performance_status
            FROM orders
            WHERE company_id = ?
            ORDER BY order_date DESC
            LIMIT 8
            """,
            (company_id,),
        )

        invoices = fetch_all(
            conn,
            """
            SELECT invoice_no, invoice_type, issue_date, amount_cny, tax_rate_pct, buyer_name
            FROM invoices
            WHERE company_id = ?
            ORDER BY issue_date DESC
            LIMIT 8
            """,
            (company_id,),
        )

        credit_history = fetch_all(
            conn,
            """
            SELECT subject_type, subject_name, credit_channel, account_count, outstanding_balance_cny,
                   overdue_count, max_overdue_bucket, hard_inquiry_3m, hard_inquiry_6m, credit_assessment, summary_note
            FROM credit_history_summaries
            WHERE company_id = ?
            ORDER BY subject_type, subject_name
            LIMIT 8
            """,
            (company_id,),
        )

        related_transactions = fetch_all(
            conn,
            """
            SELECT related_party_name, relation_type, transaction_type, transaction_amount_cny,
                   revenue_or_cost_ratio, pricing_comment, risk_level
            FROM related_party_transaction_summaries
            WHERE company_id = ?
            ORDER BY COALESCE(revenue_or_cost_ratio, 0) DESC
            LIMIT 8
            """,
            (company_id,),
        )

        risk_scores = fetch_all(
            conn,
            """
            SELECT score_dimension, score_value, score_band, rationale
            FROM case_risk_scores
            WHERE company_id = ?
            ORDER BY score_value DESC
            LIMIT 10
            """,
            (company_id,),
        )

        ip_assets = fetch_all(
            conn,
            """
            SELECT asset_type, asset_name, registration_no, ownership_holder, grant_date, status
            FROM ip_assets
            WHERE company_id = ?
            ORDER BY grant_date DESC
            LIMIT 8
            """,
            (company_id,),
        )

        innovation_qualifications = fetch_all(
            conn,
            """
            SELECT qualification_name, issuing_authority, valid_from, valid_to, status
            FROM innovation_qualifications
            WHERE company_id = ?
            ORDER BY valid_to DESC
            LIMIT 8
            """,
            (company_id,),
        )

        rd_projects = fetch_all(
            conn,
            """
            SELECT project_name, stage, budget_cny, spent_cny, expected_commercialization_date, summary
            FROM rd_projects
            WHERE company_id = ?
            ORDER BY id DESC
            LIMIT 8
            """,
            (company_id,),
        )

        public_risks = fetch_all(
            conn,
            """
            SELECT event_type, severity, event_date, title, summary, public_source_name
            FROM public_risk_events
            WHERE company_id = ?
            ORDER BY event_date DESC
            LIMIT 8
            """,
            (company_id,),
        )

        peer_comparisons = fetch_all(
            conn,
            """
            SELECT company_value, variance_pct, percentile_bucket, narrative
            FROM peer_comparisons
            WHERE company_id = ?
            ORDER BY id DESC
            LIMIT 8
            """,
            (company_id,),
        )

        industry_profile = fetch_one(
            conn,
            """
            SELECT subindustry, industry_category, policy_direction, lifecycle_stage, benchmark_note, update_cycle
            FROM industry_profiles
            WHERE subindustry = ?
            LIMIT 1
            """,
            (company["subindustry"],),
        )

    return {
        "company": company,
        "case_info": case_info,
        "recommendation": recommendation,
        "people": people,
        "related_companies": related_companies,
        "shareholding_changes": shareholding_changes,
        "profile_attributes": profile_attributes,
        "validation_findings": findings,
        "report_sections": report_sections,
        "latest_financial_period": latest_financial_period,
        "latest_metrics": latest_metrics,
        "recent_metric_history": recent_metric_history,
        "reconciliation_checks": reconciliation_checks,
        "tax_invoice_checks": tax_invoice_checks,
        "tax_filings": tax_filings,
        "bank_summaries": bank_summaries,
        "receivables": receivables,
        "social_security": social_security,
        "utilities": utilities,
        "logistics": logistics,
        "guarantees": guarantees,
        "contracts": contracts,
        "orders": orders,
        "invoices": invoices,
        "credit_history": credit_history,
        "related_transactions": related_transactions,
        "risk_scores": risk_scores,
        "ip_assets": ip_assets,
        "innovation_qualifications": innovation_qualifications,
        "rd_projects": rd_projects,
        "public_risks": public_risks,
        "peer_comparisons": peer_comparisons,
        "industry_profile": industry_profile,
        "workbook": {
            "sheet_counts": workbook_company.get("sheet_counts", {}),
            "samples": workbook_company.get("samples", {}),
            "sheet_names": workbook_index["sheet_names"],
            "overview": workbook_index["overview"],
        },
        "generated_at": now_iso(),
    }


def get_company_detail_payload(company_code: str) -> dict | None:
    detail = get_company_detail_base_payload(company_code)
    if detail is None:
        return None
    detail = supplement_poc_detail(detail)

    runtime = ensure_report_runtime(company_code, detail)
    detail["report_versions"] = [
        serialize_report_version(company_code, version) for version in runtime["versions"]
    ]
    skill_binding = get_skill_pipeline_binding(company_code)
    deepseek_configured = bool(get_deepseek_api_key())
    detail["ai_capability"] = {
        "provider": "Skill Pipeline + DeepSeek" if skill_binding else "DeepSeek",
        "configured": deepseek_configured if skill_binding else deepseek_configured,
        "mode": "skill_pipeline_via_deepseek" if skill_binding and deepseek_configured else "skill_pipeline_fallback" if skill_binding else "server_env",
        "model": f"auditing-data -> writing-report -> {DEEPSEEK_MODEL}" if skill_binding and deepseek_configured else "auditing-data -> writing-report (fallback)" if skill_binding else DEEPSEEK_MODEL,
        "generate_endpoint": f"/api/company/{quote(company_code)}/report-versions",
        "review_endpoint": f"/api/company/{quote(company_code)}/report-versions/{{version_id}}/review",
    }
    detail["skill_pipeline"] = skill_binding
    detail["due_diligence_materials"] = serialize_material_files(company_code)
    return detail


def create_report_version(company_code: str, knowledge_file_ids: list[str] | None = None, data_source_ids: list[str] | None = None) -> dict | None:
    detail = get_company_detail_base_payload(company_code)
    if detail is None:
        return None

    detail = supplement_poc_detail(detail)
    runtime = ensure_report_runtime(company_code, detail)
    version_number = runtime["next_version"]
    knowledge_map = knowledge_file_map()
    selected_ids = knowledge_file_ids or []
    selected_files = [knowledge_map[file_id] for file_id in selected_ids if file_id in knowledge_map]
    selected_data_ids = resolve_selected_data_source_ids(data_source_ids)
    references = build_reference_catalog(detail, company_code, selected_files, selected_data_ids)
    skill_binding = get_skill_pipeline_binding(company_code)
    if skill_binding:
        sections, full_text = generate_report_with_skill_pipeline(
            company_code,
            version_number,
            detail,
            selected_files,
            selected_data_ids,
        )
        sections = normalize_report_section_references(sections, references)
        full_text = normalize_reference_markers(full_text, references)
        based_on = build_skill_generation_basis(detail, company_code, selected_files, selected_data_ids)
        generation_mode = f"skill_pipeline::{DEEPSEEK_MODEL}" if get_deepseek_api_key() else "skill_pipeline::fallback"
    else:
        sections, full_text = generate_report_with_deepseek(detail, company_code, version_number, selected_files, selected_data_ids)
        sections = normalize_report_section_references(sections, references)
        full_text = normalize_reference_markers(full_text, references)
        based_on = build_generation_basis(detail, selected_files, company_code, selected_data_ids)
        generation_mode = DEEPSEEK_MODEL
    new_version = {
        "id": f"v{version_number}",
        "version_label": f"V{version_number}",
        "created_at": now_iso(),
        "status": "已生成",
        "based_on": based_on,
        "section_list": sections,
        "full_text": full_text,
        "knowledge_file_ids": [item["id"] for item in selected_files],
        "selected_data_source_ids": selected_data_ids,
        "references": references,
        "review_files": [],
        "review_result": None,
        "review_edit_text": full_text,
        "review_saved_at": None,
        "generation_mode": generation_mode,
    }
    runtime["versions"].append(new_version)
    runtime["next_version"] += 1
    return serialize_report_version(company_code, new_version)


def add_review_file(company_code: str, version_id: str, category: str, field: FieldStorage) -> dict | None:
    version = find_version(company_code, version_id)
    if version is None:
        return None

    version["review_files"].append(
        {
            "id": f"{category}-{len(version['review_files']) + 1}",
            "name": field.filename or "未命名文件",
            "category": category,
            "uploaded_at": now_iso(),
            "size": len(field.value) if field.value else 0,
            "content_excerpt": extract_upload_preview(field),
        }
    )
    version["status"] = "待预审"
    version["review_result"] = None
    return serialize_report_version(company_code, version)


def add_material_file(company_code: str, bucket: str, field: FieldStorage) -> dict:
    runtime = ensure_material_runtime(company_code)
    file_id = f"mat-{runtime['next_id']}"
    runtime["next_id"] += 1
    runtime["files"].append(
        {
            "id": file_id,
            "name": field.filename or "未命名文件",
            "bucket": bucket,
            "bucket_title": MATERIAL_BUCKET_TITLES.get(bucket, bucket),
            "uploaded_at": now_iso(),
            "size": len(field.value) if field.value else 0,
            "owner": "客户经理",
            "description": extract_upload_preview(field),
        }
    )
    return {
        "company_code": company_code,
        "files": runtime["files"],
    }


def add_knowledge_file(category: str, field: FieldStorage) -> dict:
    runtime = ensure_knowledge_runtime()
    file_id = f"kb-{runtime['next_id']}"
    runtime["next_id"] += 1
    suffix = (field.filename or "").split(".")[-1].lower() or "file"
    ensure_knowledge_library_dirs()
    stored_name = f"{file_id}.{suffix}"
    file_path = KNOWLEDGE_FILES_DIR / stored_name
    file_bytes = field.value if isinstance(field.value, bytes) else bytes(field.value or "", "utf-8")
    file_path.write_bytes(file_bytes)
    content_text = infer_file_text(file_path, extract_upload_preview(field))
    runtime["files"].append(
        {
            "id": file_id,
            "name": field.filename or "未命名文件",
            "category": category,
            "uploaded_at": now_iso(),
            "owner": "系统管理员",
            "description": extract_upload_preview(field),
            "format": suffix,
            "stored_name": stored_name,
            "content_text": content_text,
        }
    )
    save_knowledge_manifest(runtime["files"], runtime["next_id"])
    return serialize_knowledge_files()


def link_knowledge_files(company_code: str, version_id: str, file_ids: list[str]) -> dict | None:
    version = find_version(company_code, version_id)
    if version is None:
        return None
    knowledge_map = knowledge_file_map()
    knowledge_ids = set(knowledge_map.keys())
    version["knowledge_file_ids"] = [file_id for file_id in file_ids if file_id in knowledge_ids]
    detail = get_company_detail_base_payload(company_code)
    if detail is None:
        return None
    detail = supplement_poc_detail(detail)
    selected_files = [knowledge_map[file_id] for file_id in version["knowledge_file_ids"] if file_id in knowledge_map]
    version["references"] = build_reference_catalog(
        detail,
        company_code,
        selected_files,
        version.get("selected_data_source_ids"),
    )
    return serialize_report_version(company_code, version)


def save_report_version_draft(company_code: str, version_id: str, full_text: str, variant: str) -> dict | None:
    version = find_version(company_code, version_id)
    if version is None:
        return None
    if variant == "review":
        version["review_edit_text"] = (full_text or "").replace("\r", "").strip() or "当前版本暂无正文。"
        version["review_saved_at"] = now_iso()
        version["status"] = "已预审"
        return serialize_report_version(company_code, version)
    sections, normalized_text = parse_editor_report_text(full_text)
    version["section_list"] = sections
    version["full_text"] = normalized_text
    version["review_edit_text"] = normalized_text
    return serialize_report_version(company_code, version)


def get_knowledge_file(file_id: str) -> dict | None:
    return knowledge_file_map().get(file_id)


def run_report_review(company_code: str, version_id: str, knowledge_file_ids: list[str] | None = None, report_text: str | None = None) -> dict | None:
    detail = get_company_detail_base_payload(company_code)
    if detail is None:
        return None
    version = find_version(company_code, version_id)
    if version is None:
        return None
    if knowledge_file_ids is not None:
        knowledge_ids = set(knowledge_file_map().keys())
        version["knowledge_file_ids"] = [file_id for file_id in knowledge_file_ids if file_id in knowledge_ids]
    current_report_text = (report_text or version.get("review_edit_text") or version.get("full_text") or "").strip()
    knowledge_map = knowledge_file_map()
    selected_files = [knowledge_map[file_id] for file_id in version.get("knowledge_file_ids", []) if file_id in knowledge_map]
    try:
        review_result = build_ai_review_result(detail, company_code, version, current_report_text, selected_files)
    except Exception:
        review_result = build_review_result(detail, version)
    version["review_result"] = review_result
    version["status"] = "待修订"
    version["review_saved_at"] = None
    return serialize_review_result(review_result)


class DemoRequestHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def end_headers(self) -> None:
        self.send_header("Cache-Control", "no-store")
        super().end_headers()

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path

        if path == "/api/meta":
            return self.respond_json(
                {
                    "provider": "DeepSeek",
                    "deepseek_configured": bool(get_deepseek_api_key()),
                    "deepseek_model": DEEPSEEK_MODEL,
                    "pdf_export_enabled": REPORTLAB_AVAILABLE,
                    "material_buckets": [
                        {"id": key, "title": value}
                        for key, value in MATERIAL_BUCKET_TITLES.items()
                    ],
                }
            )
        if path == "/api/knowledge-files":
            return self.respond_json(serialize_knowledge_files())
        if path.startswith("/api/knowledge-files/") and path.endswith("/download"):
            file_id = unquote(path.replace("/api/knowledge-files/", "", 1).replace("/download", ""))
            file_item = get_knowledge_file(file_id)
            if file_item is None:
                return self.respond_json({"error": "Knowledge file not found"}, status=404)
            file_path = KNOWLEDGE_FILES_DIR / file_item["stored_name"]
            if not file_path.exists():
                return self.respond_json({"error": "Knowledge file binary missing"}, status=404)
            return self.respond_bytes(
                file_path.read_bytes(),
                content_type=content_type_for_suffix(file_item["format"]),
                status=200,
                filename=file_item["name"],
            )
        if path.startswith("/api/knowledge-files/") and path.endswith("/view"):
            file_id = unquote(path.replace("/api/knowledge-files/", "", 1).replace("/view", ""))
            file_item = get_knowledge_file(file_id)
            if file_item is None:
                return self.respond_json({"error": "Knowledge file not found"}, status=404)
            return self.respond_bytes(
                render_preview_html(file_item["name"], file_item["description"], file_item.get("content_text", "")),
                content_type="text/html; charset=utf-8",
                status=200,
            )
        if path == "/api/system-admin":
            return self.respond_json(build_system_admin_payload())
        if path == "/api/companies":
            return self.respond_json(get_companies_payload()["companies"])
        if path == "/api/home-feed":
            return self.respond_json(build_home_feed_payload())
        if "/report-versions/" in path and path.endswith("/pdf"):
            company_code, version_id = self.parse_version_path(path, suffix="/pdf")
            if not company_code or not version_id:
                return self.respond_json({"error": "Invalid report version path"}, status=404)
            detail = get_company_detail_payload(company_code)
            version = find_version(company_code, version_id)
            if detail is None or version is None:
                return self.respond_json({"error": "Report version not found"}, status=404)
            variant = parse_qs(parsed.query).get("variant", ["generated"])[0]
            pdf_text = version.get("review_edit_text") if variant == "review" else version["full_text"]
            pdf_bytes = build_pdf_bytes(
                f"{detail['company']['name']}尽职调查报告{'-预审修订稿' if variant == 'review' else ''}",
                version["version_label"],
                pdf_text,
            )
            filename = f"{detail['company']['name']}-{version['version_label']}-{datetime.now().strftime('%Y%m%d')}.pdf"
            return self.respond_bytes(
                pdf_bytes,
                content_type="application/pdf",
                status=200,
                filename=filename,
            )
        if path.startswith("/api/company/"):
            company_code = unquote(path.replace("/api/company/", "", 1))
            if "/" in company_code:
                company_code = company_code.split("/", 1)[0]
            payload = get_company_detail_payload(company_code)
            if payload is None:
                return self.respond_json({"error": f"Company not found: {company_code}"}, status=404)
            return self.respond_json(payload)

        if path == "/":
            self.path = "/index.html"
        elif not (ROOT / path.lstrip("/")).exists():
            self.path = "/index.html"
        return super().do_GET()

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        query = parse_qs(parsed.query)

        if path.endswith("/report-versions") and path.startswith("/api/company/"):
            company_code = unquote(path.replace("/api/company/", "", 1).replace("/report-versions", ""))
            payload = self.read_json_body()
            try:
                version = create_report_version(
                    company_code,
                    payload.get("knowledge_file_ids", []),
                    payload.get("data_source_ids", []),
                )
            except Exception as error:
                return self.respond_json({"error": str(error)}, status=502)
            if version is None:
                return self.respond_json({"error": f"Company not found: {company_code}"}, status=404)
            return self.respond_json(version, status=201)

        if path.endswith("/review") and "/report-versions/" in path:
            company_code, version_id = self.parse_version_path(path, suffix="/review")
            if not company_code or not version_id:
                return self.respond_json({"error": "Invalid review path"}, status=404)
            body = self.read_json_body()
            payload = run_report_review(
                company_code,
                version_id,
                body.get("knowledge_file_ids"),
                body.get("report_text"),
            )
            if payload is None:
                return self.respond_json({"error": "Report version not found"}, status=404)
            return self.respond_json(payload, status=201)

        if path.endswith("/review-files") and "/report-versions/" in path:
            company_code, version_id = self.parse_version_path(path, suffix="/review-files")
            if not company_code or not version_id:
                return self.respond_json({"error": "Invalid upload path"}, status=404)
            category = query.get("category", ["external"])[0]
            if category not in {"external", "internal"}:
                return self.respond_json({"error": "Invalid review file category"}, status=400)
            form = FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                },
            )
            file_field = form["file"] if "file" in form else None
            if not isinstance(file_field, FieldStorage):
                return self.respond_json({"error": "Missing file field"}, status=400)
            payload = add_review_file(company_code, version_id, category, file_field)
            if payload is None:
                return self.respond_json({"error": "Report version not found"}, status=404)
            return self.respond_json(payload, status=201)

        if path.endswith("/knowledge-links") and "/report-versions/" in path:
            company_code, version_id = self.parse_version_path(path, suffix="/knowledge-links")
            if not company_code or not version_id:
                return self.respond_json({"error": "Invalid knowledge link path"}, status=404)
            payload = self.read_json_body()
            linked = link_knowledge_files(company_code, version_id, payload.get("file_ids", []))
            if linked is None:
                return self.respond_json({"error": "Report version not found"}, status=404)
            return self.respond_json(linked, status=201)

        if path.endswith("/draft") and "/report-versions/" in path:
            company_code, version_id = self.parse_version_path(path, suffix="/draft")
            if not company_code or not version_id:
                return self.respond_json({"error": "Invalid draft path"}, status=404)
            payload = self.read_json_body()
            updated = save_report_version_draft(
                company_code,
                version_id,
                payload.get("full_text", ""),
                payload.get("variant", "generated"),
            )
            if updated is None:
                return self.respond_json({"error": "Report version not found"}, status=404)
            return self.respond_json(updated, status=201)

        if path == "/api/knowledge-files":
            category = query.get("category", ["laws"])[0]
            if category not in KNOWLEDGE_CATEGORY_TITLES:
                return self.respond_json({"error": "Invalid knowledge file category"}, status=400)
            form = FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                },
            )
            file_field = form["file"] if "file" in form else None
            if not isinstance(file_field, FieldStorage):
                return self.respond_json({"error": "Missing file field"}, status=400)
            return self.respond_json(add_knowledge_file(category, file_field), status=201)

        if path.endswith("/due-diligence-materials") and path.startswith("/api/company/"):
            company_code = unquote(path.replace("/api/company/", "", 1).replace("/due-diligence-materials", ""))
            bucket = query.get("bucket", ["application"])[0]
            if bucket not in MATERIAL_BUCKET_TITLES:
                return self.respond_json({"error": "Invalid material bucket"}, status=400)
            if get_company_detail_base_payload(company_code) is None:
                return self.respond_json({"error": f"Company not found: {company_code}"}, status=404)
            form = FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                },
            )
            file_field = form["file"] if "file" in form else None
            if not isinstance(file_field, FieldStorage):
                return self.respond_json({"error": "Missing file field"}, status=400)
            return self.respond_json(add_material_file(company_code, bucket, file_field), status=201)

        if path == "/api/deepseek/test":
            payload = self.read_json_body()
            try:
                result = call_deepseek_chat(
                    [{"role": "user", "content": payload.get("prompt", "请返回“DeepSeek连接正常”。")}],
                    system_prompt="你是银行信贷尽调智能体，请用简洁中文回答。",
                )
            except Exception as error:
                return self.respond_json({"error": str(error)}, status=502)
            return self.respond_json(result, status=201)

        return self.respond_json({"error": "Unsupported POST path"}, status=404)

    def log_message(self, fmt: str, *args) -> None:
        print(f"[demo-server] {self.address_string()} - {fmt % args}")

    def parse_version_path(self, path: str, suffix: str) -> tuple[str | None, str | None]:
        base = path[: -len(suffix)]
        if not base.startswith("/api/company/") or "/report-versions/" not in base:
            return None, None
        company_part, version_part = base.replace("/api/company/", "", 1).split("/report-versions/", 1)
        return unquote(company_part), unquote(version_part)

    def respond_json(self, payload: dict | list, status: int = 200) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def read_json_body(self) -> dict:
        content_length = int(self.headers.get("Content-Length", "0") or "0")
        if content_length <= 0:
            return {}
        raw_body = self.rfile.read(content_length)
        if not raw_body:
            return {}
        return json.loads(raw_body.decode("utf-8"))

    def respond_bytes(
        self,
        body: bytes,
        *,
        content_type: str,
        status: int = 200,
        filename: str | None = None,
    ) -> None:
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        if filename:
            self.send_header(
                "Content-Disposition",
                f"attachment; filename=report.pdf; filename*=UTF-8''{quote(filename)}",
            )
        self.end_headers()
        self.wfile.write(body)


def main() -> None:
    server = ThreadingHTTPServer((HOST, PORT), DemoRequestHandler)
    print(f"Due diligence demo server running at http://{HOST}:{PORT}")
    print(f"SQLite: {DB_PATH}")
    print(f"Excel:  {XLSX_PATH}")
    server.serve_forever()


if __name__ == "__main__":
    main()
